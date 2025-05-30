# backend/services/paper_analyzer/__init__.py
from .ai_manager import AIManager, AdvancedMemoryManager, IntentAnalysisConfig
from .paper_processor import PaperProcessor
from .rag_retriever import RAGRetriever, IndexConfig, RetrieversConfig
from .translator import Translator
from sqlalchemy.orm import Session
from models.paper import PaperAnalysis
import uuid
from sqlalchemy import func
from docx import Document
from fpdf import FPDF
import markdown
from io import BytesIO
import tempfile
import os
from datetime import datetime
from typing import List, Dict, Optional, Union
from auth.models import ChatSession, ChatMessage, SessionDocument
from models.chat import ChatFile, ChatTopic, ChatTopicSession
from contextlib import contextmanager
import time
import json
from celery import Celery
import logging
import asyncio
from dataclasses import dataclass
import psutil

logger = logging.getLogger(__name__)

@dataclass
class PaperAnalyzerConfig:
    """CargoPPT文档分析服务配置"""
    # AI配置
    enable_memory: bool = True
    enable_enhanced_intent: bool = True
    ai_temperature: float = 0.1
    max_output_tokens: int = 8192
    
    # RAG配置
    chunk_size: int = 512
    chunk_overlap: int = 50
    similarity_threshold: float = 0.7
    top_k: int = 5
    
    # 系统配置
    max_documents_per_session: int = 10
    enable_async_indexing: bool = True
    enable_caching: bool = True
    cache_ttl: int = 3600
    
    # 记忆配置
    memory_short_term_size: int = 50
    memory_long_term_size: int = 1000
    memory_importance_threshold: float = 0.7

class PaperAnalyzerService:
    def __init__(self, db: Session, config: PaperAnalyzerConfig = None):
        self.db = db
        self.config = config or PaperAnalyzerConfig()
        
        logger.info(f"[SERVICE_INIT] 初始化CargoPPT文档分析服务")
        logger.debug(f"[SERVICE_CONFIG] 配置: 记忆={self.config.enable_memory}, 增强意图={self.config.enable_enhanced_intent}")
        
        # 初始化增强的AI管理器
        self.ai_manager = AIManager(
            enable_memory=self.config.enable_memory,
            enable_enhanced_intent=self.config.enable_enhanced_intent
        )
        
        # 初始化文档处理器
        self.paper_processor = PaperProcessor()
        
        # 初始化增强的RAG检索器
        index_config = IndexConfig(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            similarity_threshold=self.config.similarity_threshold
        )
        
        retriever_config = RetrieversConfig(
            top_k=self.config.top_k,
            relevance_threshold=self.config.similarity_threshold
        )
        
        self.rag_retriever = RAGRetriever(
            db=db,
            index_config=index_config,
            retrievers_config=retriever_config
        )
        
        # 初始化翻译器
        self.translator = Translator()
        
        # 设置最大文档数限制
        self.MAX_DOCUMENTS_PER_SESSION = self.config.max_documents_per_session
        
        # 性能统计
        self.performance_stats = {
            "total_documents_processed": 0,
            "total_questions_answered": 0,
            "average_response_time": 0.0,
            "memory_usage_mb": 0.0
        }
        
        logger.info(f"[SERVICE_READY] 文档分析服务初始化完成 - 最大文档数: {self.MAX_DOCUMENTS_PER_SESSION}")

    async def analyze_paper(self, file_content: bytes, filename: str, user_id: str = None):
        """增强的文档分析 - 支持记忆存储"""
        start_time = time.time()
        try:
            logger.info(f"[ANALYZE_START] 开始分析文档 - 文件名: {filename}, 大小: {len(file_content)} bytes")
            
            # 1. 处理文档内容
            processed_result = await self.paper_processor.process(file_content, filename)
            
            # 2. 生成唯一的paper_id
            paper_id = str(uuid.uuid4())
            
            # 3. 保存到数据库
            paper_analysis = PaperAnalysis(
                paper_id=uuid.UUID(paper_id),
                filename=filename,
                content=processed_result["content"],
                line_mapping=processed_result["line_mapping"],
                total_lines=processed_result["total_lines"],
                is_scanned=any(line.get("is_scanned", False) 
                              for line in processed_result["line_mapping"].values()),
                structured_data=processed_result.get("structured_data"),
                documents=None,
                embeddings=None,
                index_built=False,
                user_id=uuid.UUID(user_id) if user_id else None  # 添加用户ID支持
            )
            self.db.add(paper_analysis)
            self.db.commit()
            
            # 4. 异步构建向量索引
            try:
                if self.config.enable_async_indexing:
                    # 异步构建索引
                    await self.rag_retriever.schedule_index_building(
                        processed_result["content"], 
                        paper_id,
                        processed_result["line_mapping"],
                        self.db,
                        processed_result.get("structured_data")
                    )
                    logger.info(f"[ANALYZE_INDEX] 已安排异步索引构建")
                else:
                    # 同步构建索引
                    await self.rag_retriever.build_index(
                        processed_result["content"], 
                        paper_id,
                        processed_result["line_mapping"],
                        self.db,
                        processed_result.get("structured_data")
                    )
                    paper_analysis.index_built = True
                    self.db.commit()
                    logger.info(f"[ANALYZE_INDEX] 索引构建完成")
            except Exception as e:
                logger.error(f"[ANALYZE_INDEX_ERROR] 索引构建失败: {str(e)}")
                # 不中断流程，允许用户先查看文档内容
            
            # 5. 记忆存储 - 存储文档信息到AI记忆中
            if self.ai_manager.memory_manager and user_id:
                document_summary = processed_result["content"][:500]  # 文档摘要
                self.ai_manager.memory_manager.add_memory(
                    content=f"用户上传文档: {filename}，内容摘要: {document_summary}",
                    context_type="document",
                    importance=0.8,
                    user_id=user_id,
                    document_id=paper_id
                )
                logger.debug(f"[ANALYZE_MEMORY] 已存储文档到记忆系统")
            
            # 6. 更新性能统计
            processing_time = time.time() - start_time
            self.performance_stats["total_documents_processed"] += 1
            self.performance_stats["average_response_time"] = (
                self.performance_stats["average_response_time"] * 0.9 + processing_time * 0.1
            )
            
            logger.info(f"[ANALYZE_SUCCESS] 文档分析完成 - 耗时: {processing_time:.2f}s")
            
            return {
                "status": "success",
                "message": "文档分析完成",
                "paper_id": paper_id,
                "content": processed_result["content"],
                "line_mapping": processed_result["line_mapping"] or {},
                "total_lines": processed_result["total_lines"],
                "is_scanned": paper_analysis.is_scanned,
                "has_structured_data": "structured_data" in processed_result  # 指示是否包含结构化数据
            }
        except Exception as e:
            self.db.rollback()
            print(f"Paper analysis error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

    async def get_paper_content(self, paper_id: str) -> str:
        """获取文档内容"""
        try:
            paper = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.paper_id == uuid.UUID(paper_id)
            ).first()
            
            if not paper:
                raise Exception("Paper not found")
                
            return paper.content
        except Exception as e:
            raise Exception(f"Failed to get paper content: {str(e)}")

    async def ask_question(self, question: str, paper_id: str = None, context: str = None, session_id: str = None, user_id: str = None):
        """向文档提问，仅支持会话模式"""
        try:
            # 必须提供会话ID和用户ID
            if not session_id or not user_id:
                return {
                    "status": "error",
                    "message": "必须提供会话ID和用户ID，不再支持单次问答模式"
                }
            
            # 使用会话模式处理
            return await self.send_message(
                session_id=session_id,
                message=question,
                user_id=user_id
            )
            
        except Exception as e:
            return {
                "status": "error", 
                "message": str(e)
            }

    async def get_question_history(self, paper_id: str):
        """获取问答历史 - 已废弃，请使用会话模式"""
        raise Exception("单次问答模式已废弃，请使用会话模式获取历史消息")

    async def translate_paper(self, paper_id: str, target_lang: str):
        """翻译文档内容"""
        try:
            # 获取文档内容
            paper = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.paper_id == uuid.UUID(paper_id)
            ).first()
            
            if not paper:
                raise Exception("Paper not found")
            
            # 验证目标语言
            if target_lang not in self.translator.supported_languages:
                raise ValueError(f"Unsupported target language: {target_lang}")
            
            # 翻译内容
            translated_content = await self.translator.translate_text(
                paper.content,
                target_lang
            )
            
            # 更新数据库中的翻译内容
            paper.translated_content = translated_content
            paper.translation_lang = target_lang
            paper.translation_updated_at = func.now()
            self.db.commit()
            
            return {
                "status": "success",
                "content": translated_content,
                "language": target_lang
            }
        except Exception as e:
            self.db.rollback()
            return {
                "status": "error",
                "message": str(e)
            }

    async def get_supported_languages(self):
        """获取支持的语言列表"""
        return self.translator.get_supported_languages()

    async def download_translation(self, paper_id: str, target_lang: str, format: str) -> bytes:
        """下载翻译结果"""
        try:
            # 获取文档内容
            paper = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.paper_id == uuid.UUID(paper_id)
            ).first()
            
            if not paper:
                raise Exception("Paper not found")
            
            # 获取翻译内容
            translated_content = await self.translator.translate_text(
                paper.content,
                target_lang
            )
            
            # 根据格式转换内容
            if format == 'docx':
                return await self._convert_to_docx(translated_content)
            elif format == 'pdf':
                return await self._convert_to_pdf(translated_content)
            elif format == 'md':
                return translated_content.encode('utf-8')
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            raise Exception(f"Download error: {str(e)}")

    async def _convert_to_docx(self, content: str) -> bytes:
        """转换为 Word 文档"""
        try:
            doc = Document()
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # 保存到内存
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()
        except Exception as e:
            raise Exception(f"Failed to convert to DOCX: {str(e)}")

    async def _convert_to_pdf(self, content: str) -> bytes:
        """转换为 PDF 文档"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # 处理内容
            for line in content.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 10, txt=line)
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_path = temp_file.name
            
            # 保存到临时文件
            pdf.output(temp_path)
            
            # 读取文件内容
            with open(temp_path, 'rb') as f:
                content = f.read()
            
            # 删除临时文件
            try:
                os.unlink(temp_path)
            except:
                pass
            
            return content
            
        except Exception as e:
            # 确保清理临时文件
            if 'temp_path' in locals():
                try:
                    os.unlink(temp_path)
                except:
                    pass
            raise Exception(f"Failed to convert to PDF: {str(e)}")

    async def create_chat_session(self, user_id: str, title: str = None, paper_ids: List[str] = None, session_type: str = "document", is_temporary: bool = False) -> dict:
        """创建新的对话会话，支持无文档和多文档模式，增加临时会话标记"""
        try:
            # Debug logs
            print(f"Creating session: user_id={user_id}, title={title}, paper_ids={paper_ids}, type={session_type}, temporary={is_temporary}")
            
            if not title:
                if session_type == "document" and paper_ids and len(paper_ids) > 0:
                    # 获取第一个文档的文件名作为标题
                    paper = self.db.query(PaperAnalysis).filter(
                        PaperAnalysis.paper_id == uuid.UUID(paper_ids[0])
                    ).first()
                    title = f"关于 {paper.filename if paper else '文档'} 的对话"
                else:
                    title = f"新对话 {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
            
            # 创建会话
            session = ChatSession(
                user_id=uuid.UUID(user_id),
                title=title,
                session_type=session_type,
                is_temporary=is_temporary,  # 添加临时标记
                # 如果是单文档模式，保持向后兼容，为paper_id赋值
                paper_id=uuid.UUID(paper_ids[0]) if session_type == "document" and paper_ids and len(paper_ids) > 0 else None
            )
            self.db.add(session)
            self.db.flush()  # 获取生成的ID
            
            # 如果有文档IDs，创建关联关系
            if paper_ids and len(paper_ids) > 0:
                # 检查文档数量是否超出限制
                if len(paper_ids) > self.MAX_DOCUMENTS_PER_SESSION:
                    raise ValueError(f"一个会话最多支持{self.MAX_DOCUMENTS_PER_SESSION}个文档")
                
                # 关联文档
                for i, paper_id in enumerate(paper_ids):
                    # 获取文档信息
                    paper = self.db.query(PaperAnalysis).filter(
                        PaperAnalysis.paper_id == uuid.UUID(paper_id)
                    ).first()
                    
                    if not paper:
                        raise ValueError(f"文档ID {paper_id} 不存在")
                    
                    # 创建关联
                    session_doc = SessionDocument(
                        session_id=session.id,
                        paper_id=uuid.UUID(paper_id),
                        order=i,
                        filename=paper.filename
                    )
                    self.db.add(session_doc)
            
            self.db.commit()
            
            # 返回会话信息
            documents = []
            if paper_ids and len(paper_ids) > 0:
                # 获取关联的文档信息
                session_docs = self.db.query(SessionDocument).filter(
                    SessionDocument.session_id == session.id
                ).order_by(SessionDocument.order).all()
                
                documents = [{
                    "id": str(doc.id),
                    "paper_id": str(doc.paper_id),
                    "filename": doc.filename,
                    "order": doc.order
                } for doc in session_docs]
            
            return {
                "id": str(session.id),
                "title": session.title,
                "created_at": session.created_at.isoformat() if session.created_at else datetime.utcnow().isoformat(),
                "updated_at": session.updated_at.isoformat() if session.updated_at else datetime.utcnow().isoformat(),
                "paper_id": str(session.paper_id) if session.paper_id else None,  # 向后兼容
                "session_type": session.session_type,
                "is_temporary": session.is_temporary,
                "message_count": 0,
                "last_message": "",
                "documents": documents
            }
        except Exception as e:
            self.db.rollback()
            print(f"Error creating chat session: {str(e)}")
            import traceback
            print(traceback.format_exc())
            raise Exception(f"Failed to create chat session: {str(e)}")

    async def get_user_chat_sessions(self, user_id: str, paper_id: str = None) -> List[dict]:
        """获取用户的所有对话会话"""
        try:
            # 基础查询
            query = self.db.query(ChatSession).filter(
                ChatSession.user_id == uuid.UUID(user_id),
                ChatSession.is_active == True
            )
            
            # 如果指定了paper_id，只获取包含该文档的会话
            if paper_id:
                # 方法1：使用传统的paper_id字段（向后兼容）
                # 方法2：通过SessionDocument关联查询
                query = query.filter(
                    (ChatSession.paper_id == uuid.UUID(paper_id)) | 
                    (ChatSession.id.in_(
                        self.db.query(SessionDocument.session_id).filter(
                            SessionDocument.paper_id == uuid.UUID(paper_id)
                        )
                    ))
                )
            
            # 按最后更新时间排序
            sessions = query.order_by(ChatSession.updated_at.desc()).all()
            
            # 构建响应
            result = []
            for session in sessions:
                # 获取最后一条消息
                last_message = self.db.query(ChatMessage).filter(
                    ChatMessage.session_id == session.id
                ).order_by(ChatMessage.created_at.desc()).first()
                
                # 获取关联的文档信息
                documents = []
                if session.session_type == "document":
                    session_docs = self.db.query(SessionDocument).filter(
                        SessionDocument.session_id == session.id
                    ).order_by(SessionDocument.order).all()
                    
                    documents = [{
                        "id": str(doc.id),
                        "paper_id": str(doc.paper_id),
                        "filename": doc.filename,
                        "order": doc.order
                    } for doc in session_docs]
                
                # 消息数量
                message_count = self.db.query(ChatMessage).filter(
                    ChatMessage.session_id == session.id
                ).count()
                
                result.append({
                    "id": str(session.id),
                    "title": session.title,
                    "created_at": session.created_at.isoformat(),
                    "updated_at": session.updated_at.isoformat(),
                    "paper_id": str(session.paper_id) if session.paper_id else None,  # 向后兼容
                    "session_type": session.session_type,
                    "message_count": message_count,
                    "last_message": last_message.content if last_message else "",
                    "documents": documents
                })
            
            return result
        except Exception as e:
            raise Exception(f"Failed to get chat sessions: {str(e)}")

    async def get_full_message_content(self, message_id: str) -> dict:
        """获取完整的消息内容，包括分片聚合"""
        try:
            # 获取主消息
            message = self.db.query(ChatMessage).filter(
                ChatMessage.id == uuid.UUID(message_id)
            ).first()
            
            if not message:
                raise ValueError(f"消息ID {message_id} 不存在")
            
            # 如果没有分片，直接返回
            if not message.has_fragments:
                return {
                    "id": str(message.id),
                    "content": message.content,
                    "rich_content": message.rich_content or [],
                    "has_fragments": False,
                    "fragment_count": 0
                }
            
            # 获取所有分片
            from models.chat import MessageFragment
            fragments = self.db.query(MessageFragment).filter(
                MessageFragment.message_id == message.id
            ).order_by(MessageFragment.fragment_index).all()
            
            if not fragments:
                # 没有找到分片，返回主消息内容
                return {
                    "id": str(message.id),
                    "content": message.content,
                    "rich_content": message.rich_content or [],
                    "has_fragments": True,
                    "fragment_count": 0,
                    "warning": "分片数据丢失"
                }
            
            # 聚合分片内容
            aggregated_content = []
            for fragment in fragments:
                try:
                    if isinstance(fragment.content, list):
                        aggregated_content.extend(fragment.content)
                    elif isinstance(fragment.content, str):
                        # 尝试解析JSON
                        import json
                        parsed_content = json.loads(fragment.content)
                        if isinstance(parsed_content, list):
                            aggregated_content.extend(parsed_content)
                        else:
                            aggregated_content.append(parsed_content)
                    else:
                        aggregated_content.append(fragment.content)
                except Exception as parse_error:
                    print(f"解析分片 {fragment.fragment_index} 失败: {str(parse_error)}")
                    # 添加错误占位符
                    aggregated_content.append({
                        "type": "markdown",
                        "content": f"⚠️ 分片 {fragment.fragment_index} 解析失败"
                    })
            
            # 验证内容完整性
            expected_fragments = len(fragments)
            actual_fragments = len([f for f in fragments if f.content])
            
            result = {
                "id": str(message.id),
                "content": message.content,  # 保留原始摘要内容
                "rich_content": aggregated_content,
                "has_fragments": True,
                "fragment_count": len(fragments),
                "integrity_check": {
                    "expected_fragments": expected_fragments,
                    "actual_fragments": actual_fragments,
                    "is_complete": expected_fragments == actual_fragments
                }
            }
            
            if not result["integrity_check"]["is_complete"]:
                result["warning"] = f"内容可能不完整：期望{expected_fragments}个分片，实际{actual_fragments}个"
            
            return result
            
        except Exception as e:
            print(f"获取完整消息内容失败: {str(e)}")
            raise Exception(f"获取完整消息内容失败: {str(e)}")

    async def get_chat_history(self, session_id: str, limit: int = 20, before_id: Optional[str] = None) -> dict:
        """获取会话历史消息，自动处理分片聚合"""
        try:
            # 简化错误处理，直接使用try/except包裹整个逻辑
            session_uuid = uuid.UUID(session_id)
            
            # 简化查询，直接获取消息 - 修复：排除流式状态的消息
            query = self.db.query(ChatMessage).filter(
                ChatMessage.session_id == session_uuid,
                ChatMessage.message_type != 'streaming'  # ✅ 排除流式状态消息
            )
            
            # 处理before_id参数
            if before_id:
                try:
                    before_uuid = uuid.UUID(before_id)
                    before_message = self.db.query(ChatMessage).filter(
                        ChatMessage.id == before_uuid
                    ).first()
                    if before_message:
                        query = query.filter(
                            ChatMessage.created_at < before_message.created_at
                        )
                except Exception:
                    # 忽略无效的before_id
                    pass
            
            # 获取按时间排序的消息 - 按创建时间正序排列
            messages = query.order_by(ChatMessage.created_at.asc()).limit(limit).all()
            
            # 构建简化的响应数据
            result = []
            for message in messages:
                try:
                    # 检查是否有分片，如果有则聚合
                    if message.has_fragments:
                        try:
                            full_content = await self.get_full_message_content(str(message.id))
                            rich_content = full_content.get("rich_content", [])
                            
                            # 添加分片信息到消息中
                            if full_content.get("warning"):
                                rich_content.append({
                                    "type": "markdown",
                                    "content": f"⚠️ {full_content['warning']}"
                                })
                        except Exception as fragment_error:
                            print(f"聚合分片失败，使用原始内容: {str(fragment_error)}")
                            rich_content = message.rich_content or [{"type": "markdown", "content": message.content or ""}]
                    else:
                        rich_content = message.rich_content or [{"type": "markdown", "content": message.content or ""}]
                    
                    # 避免任何可能的空值或类型错误
                    message_data = {
                "id": str(message.id),
                        "role": message.role or "",
                        "content": message.content or "",
                        "created_at": message.created_at.isoformat() if message.created_at else "",
                        "sources": message.sources if isinstance(message.sources, list) else [],
                        "confidence": float(message.confidence) if message.confidence is not None else 0.0,
                        "reply": rich_content,
                        "has_fragments": bool(message.has_fragments)
                    }
                    
                    result.append(message_data)
                except Exception as msg_error:
                    # 如果单个消息处理失败，记录错误但继续处理其他消息
                    print(f"Error processing message {message.id}: {str(msg_error)}")
                    continue
            
            return {
                "messages": result,
                "has_more": len(messages) >= limit
            }
            
        except Exception as e:
            # 记录详细错误，但返回空结果而不是抛出异常
            import traceback
            print(f"Error in get_chat_history: {str(e)}")
            print(traceback.format_exc())
            return {
                "messages": [],
                "has_more": False
            }

    async def send_message(self, session_id: str, message: str, user_id: str) -> dict:
        """发送消息到特定会话 - 集成记忆管理和增强意图分析"""
        # 添加消息状态跟踪
        message_status = {"stored": False, "user_message_id": None, "ai_message_id": None}
        user_message = None
        user_message_stored = False
        
        # 记录开始处理日志
        logger.info(f"[MESSAGE_START] 开始处理消息 - 会话ID: {session_id}, 用户ID: {user_id}, 消息长度: {len(message)} 字符")
        logger.debug(f"[MESSAGE_CONTENT] 消息内容预览: {message[:100]}{'...' if len(message) > 100 else ''}")
        
        # 第一个事务：存储用户消息
        try:
            logger.info(f"[USER_MESSAGE_SAVE] 开始保存用户消息 - 会话ID: {session_id}")
            current_time = datetime.utcnow()
            
            user_message = ChatMessage(
                session_id=uuid.UUID(session_id),
                role='user',
                content=message,
                message_type='markdown',
                created_at=current_time
            )
            self.db.add(user_message)
            self.db.commit()
            user_message_stored = True
            message_status["user_message_id"] = str(user_message.id)
            logger.info(f"[USER_MESSAGE_SAVED] 用户消息保存成功 - 消息ID: {user_message.id}")
        except Exception as user_msg_error:
            self.db.rollback()
            logger.error(f"[USER_MESSAGE_ERROR] 用户消息保存失败 - 会话ID: {session_id}, 错误: {str(user_msg_error)}")
            # 继续尝试处理
        
        # 获取会话和生成响应 - 不在事务中进行
        session = None
        response = None
        processing_errors = []
        context_history = ""
        
        try:
            logger.info(f"[SESSION_FETCH] 开始获取会话信息 - 会话ID: {session_id}")
            # 获取会话信息
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                error_msg = "Chat session not found"
                logger.error(f"[SESSION_ERROR] 会话不存在 - 会话ID: {session_id}, 用户ID: {user_id}")
                return {
                    "user_message": self._format_message(user_message),
                    "ai_message": self._format_ai_message({
                        "answer": error_msg,
                        "sources": [],
                        "confidence": 0.0,
                        "reply": [{"type": "markdown", "content": error_msg}]
                    })
                }
            
            logger.info(f"[SESSION_FOUND] 会话信息获取成功 - 标题: {session.title}, 类型: {session.session_type}")
            
            # 获取历史消息作为上下文 - 独立查询
            try:
                logger.info(f"[HISTORY_FETCH] 开始获取历史消息 - 会话ID: {session_id}")
                history_messages = self.db.query(ChatMessage).filter(
                    ChatMessage.session_id == uuid.UUID(session_id)
                ).order_by(ChatMessage.created_at.asc()).all()
                
                logger.info(f"[HISTORY_FOUND] 历史消息获取成功 - 总数: {len(history_messages)} 条")
                
                # 构建上下文历史
                context_history = "\n".join([
                    f"{'User' if msg.role == 'user' else 'Assistant'}: {msg.content}"
                    for msg in history_messages[-8:]
                ])
                
                logger.debug(f"[CONTEXT_BUILT] 上下文构建完成 - 使用最近 {min(8, len(history_messages))} 条消息")
            except Exception as history_error:
                logger.error(f"[HISTORY_ERROR] 获取历史消息失败 - 会话ID: {session_id}, 错误: {str(history_error)}")
                context_history = ""  # 使用空历史
            
            # 获取会话相关文档
            paper_ids = []
            has_documents = False
            document_info = {"count": 0, "names": []}
            
            if session.session_type != "general":
                logger.info(f"[DOCUMENTS_FETCH] 开始获取会话文档 - 会话ID: {session_id}")
                # 获取关联文档ID
                session_docs = self.db.query(SessionDocument).filter(
                    SessionDocument.session_id == uuid.UUID(session_id)
                ).all()
                
                for doc in session_docs:
                    paper_ids.append(str(doc.paper_id))
                    document_info["names"].append(doc.filename)

                has_documents = len(paper_ids) > 0
                document_info["count"] = len(paper_ids)
                logger.info(f"[DOCUMENTS_FOUND] 会话文档获取完成 - 文档数量: {len(paper_ids)}")
                if paper_ids:
                    logger.debug(f"[DOCUMENTS_LIST] 文档ID列表: {paper_ids}")
                    logger.debug(f"[DOCUMENTS_DETAIL] 文档详情: {document_info}")
            else:
                logger.info(f"[DOCUMENTS_SKIP] 通用会话，跳过文档获取")
            
            # 将历史消息转换为结构化格式
            chat_history = []
            for msg in history_messages[-8:]:
                chat_history.append({
                    "role": "user" if msg.role == "user" else "assistant",
                    "content": msg.content
                })

            # 智能识别用户意图 - 使用简化的意图识别
            logger.info(f"[INTENT_ANALYSIS] 开始简化意图分析 - 消息: {message[:50]}...")
            intent_result = await self.ai_manager.identify_intent(
                question=message,
                context=None,
                has_documents=has_documents,
                history=context_history,
                document_info=document_info
            )

            logger.info(f"[INTENT_RESULT] 意图分析完成 - 结果: {intent_result}, 有文档: {has_documents}")

            # 根据意图选择处理模式 - 简化的二选一路由
            if intent_result == "GENERAL_QUERY":
                # 通用查询模式
                logger.info(f"[AI_MODE] 使用通用聊天模式处理 - 会话ID: {session_id}")
                response = await self.ai_manager.chat_without_context(
                    message=message,
                    history=chat_history
                )
                logger.info(f"[AI_RESPONSE] 通用模式AI响应完成 - 响应长度: {len(response.get('answer', ''))} 字符")
            
            elif intent_result == "DOCUMENT_QUERY":
                # 文档查询模式
                logger.info(f"[AI_MODE] 使用文档模式处理 - 会话ID: {session_id}")
                retrieval_context = None

                # 获取文档上下文
                if paper_ids:
                    try:
                        logger.info(f"[CONTEXT_RETRIEVAL] 开始文档上下文检索 - 文档数量: {len(paper_ids)}")
                        
                        # 使用标准检索参数
                        doc_limit_per_paper = 4
                        
                        if len(paper_ids) > 1:
                            logger.info(f"[CONTEXT_MULTI] 多文档上下文检索 - 每文档限制: {doc_limit_per_paper}")
                            retrieval_context = await self.rag_retriever.get_context_from_multiple_docs(
                                question=message,
                                paper_ids=paper_ids,
                                history=context_history,
                                doc_limit_per_paper=doc_limit_per_paper
                            )
                        else:
                            logger.info(f"[CONTEXT_SINGLE] 单文档上下文检索 - 文档ID: {paper_ids[0]}")
                            retrieval_context = await self.rag_retriever.get_relevant_context(
                                question=message,
                                paper_id=paper_ids[0],
                                history=context_history,
                                doc_limit=doc_limit_per_paper * 2
                            )
                        
                        if retrieval_context:
                            context_info = self._analyze_retrieval_context(retrieval_context)
                            logger.info(f"[CONTEXT_SUCCESS] 文档上下文获取成功 - {context_info}")
                        else:
                            logger.warning(f"[CONTEXT_EMPTY] 未检索到相关文档内容")
                        
                        logger.info(f"[AI_DOCUMENT] 开始文档模式AI处理")
                        response = await self.ai_manager.get_response(
                            question=message,
                            context=retrieval_context,
                            history=context_history
                        )
                        logger.info(f"[AI_RESPONSE] 文档模式AI响应完成 - 响应长度: {len(response.get('answer', ''))} 字符, 来源数量: {len(response.get('sources', []))}")
                        
                    except Exception as context_error:
                        logger.error(f"[CONTEXT_ERROR] 获取文档上下文失败，降级为通用模式 - 错误: {str(context_error)}")
                        # 降级为通用模式
                        logger.info(f"[AI_FALLBACK] 文档检索失败，降级为通用模式")
                        response = await self.ai_manager.chat_without_context(
                            message=message,
                            history=chat_history
                        )
                        # 在响应中添加错误说明
                        original_answer = response.get("answer", "")
                        error_note = f"注意：无法访问文档内容（{str(context_error)[:50]}），以下回答基于通用知识。\n\n"
                        response["answer"] = error_note + original_answer
                        logger.info(f"[AI_FALLBACK] 降级通用模式响应完成 - 响应长度: {len(response.get('answer', ''))} 字符")
                else:
                    # 没有文档时降级为通用模式
                    logger.info(f"[AI_FALLBACK] 无文档，降级为通用模式")
                    response = await self.ai_manager.chat_without_context(
                        message=message,
                        history=chat_history
                    )
                    logger.info(f"[AI_RESPONSE] 无文档降级通用模式响应完成 - 响应长度: {len(response.get('answer', ''))} 字符")
            
            else:
                # 未知意图或fallback，使用通用模式
                logger.warning(f"[AI_MODE] 未知意图: {intent_result}, 使用通用模式")
                response = await self.ai_manager.chat_without_context(
                    message=message,
                    history=chat_history
                )
                logger.info(f"[AI_RESPONSE] 未知意图通用模式响应完成 - 响应长度: {len(response.get('answer', ''))} 字符")
            
        except Exception as process_error:
            logger.error(f"[PROCESSING_ERROR] 消息处理过程中发生错误 - 会话ID: {session_id}, 错误: {str(process_error)}")
            import traceback
            logger.debug(f"[PROCESSING_TRACEBACK] 错误堆栈: {traceback.format_exc()}")
            response = {
                "answer": f"处理消息时发生错误: {str(process_error)}",
                "sources": [],
                "confidence": 0.0,
                "reply": [{"type": "markdown", "content": f"处理消息时发生错误: {str(process_error)}"}]
            }
        
        # 第三个事务：存储AI响应
        ai_message = None
        try:
            logger.info(f"[AI_SAVE] 开始保存AI响应 - 会话ID: {session_id}")
            # 使用独立事务存储AI响应
            if response:
                logger.debug(f"[AI_CONTENT] AI响应内容预览: {response.get('answer', '')[:100]}{'...' if len(response.get('answer', '')) > 100 else ''}")
                ai_message = await self._save_ai_response_with_retry(session_id, response)
                if ai_message:
                    message_status["ai_message_id"] = str(ai_message.id)
                    logger.info(f"[AI_SAVED] AI响应保存成功 - 消息ID: {ai_message.id}")
                    
                    # 更新会话时间 - 轻量级操作
                    if session:
                        session.updated_at = datetime.utcnow()
                        self.db.commit()
                        logger.debug(f"[SESSION_UPDATED] 会话更新时间已更新")
                    message_status["stored"] = True
                else:
                    logger.error(f"[AI_SAVE_FAILED] AI响应保存失败 - 返回None")
            
        except Exception as save_error:
            self.db.rollback()
            logger.error(f"[AI_SAVE_ERROR] AI响应保存异常 - 会话ID: {session_id}, 错误: {str(save_error)}")
            # 创建一个错误响应
            message_status["error"] = str(save_error)
        
        # 记录处理完成
        logger.info(f"[MESSAGE_COMPLETE] 消息处理完成 - 会话ID: {session_id}, 用户消息ID: {message_status.get('user_message_id')}, AI消息ID: {message_status.get('ai_message_id')}")
        
        # 返回结果
        return {
            "user_message": self._format_message(user_message),
            "ai_message": self._format_ai_message(response, ai_message),
            "_message_status": message_status
        }

    async def _save_message_fragments(self, message_id, content_chunks):
        """可靠的消息片段存储，使用独立事务，增强异常处理"""
        from models.chat import MessageFragment
        
        logger.info(f"[FRAGMENT_SAVE_START] 开始保存消息分片 - 消息ID: {message_id}, 分片数量: {len(content_chunks) if content_chunks else 0}")
        
        successful = False
        retries = 0
        max_retries = 3
        
        while not successful and retries < max_retries:
            try:
                logger.info(f"[FRAGMENT_SAVE_ATTEMPT] 第 {retries + 1} 次分片保存尝试 - 消息ID: {message_id}")
                
                # 使用独立事务会话
                from sqlalchemy.orm import sessionmaker
                from database import engine
                
                SessionLocal = sessionmaker(bind=engine)
                fragment_session = SessionLocal()
                
                try:
                    logger.debug(f"[FRAGMENT_CLEANUP] 清除旧分片 - 消息ID: {message_id}")
                    # 先清除可能存在的旧片段
                    deleted_count = fragment_session.query(MessageFragment).filter(
                        MessageFragment.message_id == message_id
                    ).delete()
                    logger.debug(f"[FRAGMENT_CLEANUP_DONE] 清除了 {deleted_count} 个旧分片")
                    
                    # 验证和清理content_chunks
                    logger.info(f"[FRAGMENT_VALIDATE] 开始验证内容块 - 原始数量: {len(content_chunks) if content_chunks else 0}")
                    validated_chunks = self._validate_content_chunks(content_chunks)
                    logger.info(f"[FRAGMENT_VALIDATE_DONE] 内容块验证完成 - 有效数量: {len(validated_chunks)}")
                    
                    if not validated_chunks:
                        logger.warning(f"[FRAGMENT_NO_VALID] 没有有效的内容块可以存储 - 消息ID: {message_id}")
                        fragment_session.close()
                        return False
                    
                    # 批量添加新片段
                    logger.info(f"[FRAGMENT_CREATE] 开始创建分片对象 - 数量: {len(validated_chunks)}")
                    fragments = []
                    for i, chunk in enumerate(validated_chunks):
                        try:
                            logger.debug(f"[FRAGMENT_PROCESS] 处理分片 #{i} - 类型: {type(chunk)}")
                            # 安全序列化chunk
                            serialized_chunk = self._safe_serialize_chunk(chunk)
                            content_hash = self._calculate_content_hash(serialized_chunk)
                            
                            fragment = MessageFragment(
                                message_id=message_id,
                                fragment_index=i,
                                content=serialized_chunk,
                                created_at=datetime.utcnow(),
                                content_hash=content_hash
                            )
                            fragments.append(fragment)
                            logger.debug(f"[FRAGMENT_CREATED] 分片 #{i} 创建成功 - 哈希: {content_hash[:8]}...")
                        except Exception as chunk_error:
                            logger.error(f"[FRAGMENT_PROCESS_ERROR] 处理分片 #{i} 失败: {str(chunk_error)}")
                            # 创建错误占位符
                            error_chunk = {"type": "markdown", "content": f"⚠️ 分片 {i} 处理失败"}
                            serialized_error = self._safe_serialize_chunk(error_chunk)
                            
                            fragment = MessageFragment(
                                message_id=message_id,
                                fragment_index=i,
                                content=serialized_error,
                                created_at=datetime.utcnow(),
                                content_hash=self._calculate_content_hash(serialized_error)
                            )
                            fragments.append(fragment)
                            logger.debug(f"[FRAGMENT_ERROR_PLACEHOLDER] 为分片 #{i} 创建了错误占位符")
                    
                    if not fragments:
                        logger.error(f"[FRAGMENT_CREATE_FAILED] 没有成功创建任何分片 - 消息ID: {message_id}")
                        fragment_session.close()
                        return False
                    
                    logger.info(f"[FRAGMENT_BULK_SAVE] 开始批量保存分片 - 数量: {len(fragments)}")
                    fragment_session.bulk_save_objects(fragments)
                    fragment_session.commit()
                    logger.info(f"[FRAGMENT_BULK_SAVE_DONE] 分片批量保存完成")
                    
                    # 验证片段数量和内容
                    logger.info(f"[FRAGMENT_VERIFY] 开始验证保存的分片")
                    stored_fragments = fragment_session.query(MessageFragment).filter(
                        MessageFragment.message_id == message_id
                    ).order_by(MessageFragment.fragment_index).all()
                    
                    logger.info(f"[FRAGMENT_VERIFY_COUNT] 验证分片数量 - 期望: {len(fragments)}, 实际: {len(stored_fragments)}")
                    
                    if len(stored_fragments) == len(fragments):
                        # 验证内容哈希
                        logger.debug(f"[FRAGMENT_VERIFY_HASH] 开始验证分片哈希")
                        validated = True
                        for i, fragment in enumerate(stored_fragments):
                            if i < len(validated_chunks):
                                expected_chunk = validated_chunks[i]
                                expected_serialized = self._safe_serialize_chunk(expected_chunk)
                                expected_hash = self._calculate_content_hash(expected_serialized)
                                
                                if fragment.content_hash != expected_hash:
                                    logger.error(f"[FRAGMENT_HASH_MISMATCH] 分片 #{i} 哈希不匹配 - 期望: {expected_hash[:8]}..., 实际: {fragment.content_hash[:8]}...")
                                    validated = False
                                    break
                                else:
                                    logger.debug(f"[FRAGMENT_HASH_OK] 分片 #{i} 哈希验证通过")
                        
                        if validated:
                            successful = True
                            logger.info(f"[FRAGMENT_SAVE_SUCCESS] 成功存储并验证了 {len(fragments)} 个分片 - 消息ID: {message_id}")
                        else:
                            logger.error(f"[FRAGMENT_HASH_VALIDATION_FAILED] 分片哈希验证失败 - 消息ID: {message_id}")
                    else:
                        logger.error(f"[FRAGMENT_COUNT_MISMATCH] 分片数量不匹配 - 期望: {len(fragments)}, 实际: {len(stored_fragments)}")
                        
                except Exception as inner_error:
                    fragment_session.rollback()
                    logger.error(f"[FRAGMENT_TRANSACTION_ERROR] 分片事务错误: {str(inner_error)}")
                    import traceback
                    logger.debug(f"[FRAGMENT_TRANSACTION_TRACEBACK] 事务错误堆栈: {traceback.format_exc()}")
                    raise
                finally:
                    fragment_session.close()
                    logger.debug(f"[FRAGMENT_SESSION_CLOSED] 分片会话已关闭")
                
                if successful:
                    logger.info(f"[FRAGMENT_SAVE_COMPLETE] 分片保存完成 - 消息ID: {message_id}")
                    return True
                
            except Exception as outer_error:
                retries += 1
                logger.error(f"[FRAGMENT_SAVE_RETRY] 分片保存尝试 #{retries} 失败: {str(outer_error)}")
                if retries < max_retries:
                    sleep_time = 0.5 * retries
                    logger.info(f"[FRAGMENT_SAVE_WAIT] 等待 {sleep_time} 秒后重试")
                    await asyncio.sleep(sleep_time)  # 指数退避
                    continue
        
        logger.error(f"[FRAGMENT_SAVE_EXHAUSTED] 所有分片存储尝试失败 - 消息ID: {message_id}")
        return False

    def _validate_content_chunks(self, content_chunks):
        """验证和清理内容块"""
        logger.info(f"[CHUNK_VALIDATE_START] 开始验证内容块 - 输入数量: {len(content_chunks) if content_chunks else 0}")
        
        if not content_chunks:
            logger.warning(f"[CHUNK_VALIDATE_EMPTY] 输入内容块为空")
            return []
        
        validated_chunks = []
        for i, chunk in enumerate(content_chunks):
            try:
                logger.debug(f"[CHUNK_VALIDATE_ITEM] 验证内容块 #{i} - 类型: {type(chunk)}")
                
                # 验证chunk是否可序列化
                if chunk is None:
                    logger.debug(f"[CHUNK_VALIDATE_NULL] 跳过空内容块 #{i}")
                    continue
                
                # 尝试序列化测试
                test_serialized = self._safe_serialize_chunk(chunk)
                if test_serialized:
                    validated_chunks.append(chunk)
                    logger.debug(f"[CHUNK_VALIDATE_OK] 内容块 #{i} 验证通过")
                else:
                    logger.warning(f"[CHUNK_VALIDATE_SERIALIZE_FAIL] 内容块 #{i} 序列化失败，跳过")
                    
            except Exception as e:
                logger.error(f"[CHUNK_VALIDATE_ERROR] 验证内容块 #{i} 失败: {str(e)}")
                # 尝试创建安全的替代chunk
                try:
                    safe_chunk = {"type": "markdown", "content": f"⚠️ 原始内容无法处理: {str(chunk)[:100]}..."}
                    validated_chunks.append(safe_chunk)
                    logger.debug(f"[CHUNK_VALIDATE_SAFE_REPLACEMENT] 为内容块 #{i} 创建了安全替代")
                except:
                    logger.error(f"[CHUNK_VALIDATE_REPLACEMENT_FAIL] 无法为内容块 #{i} 创建安全替代")
                    continue
        
        logger.info(f"[CHUNK_VALIDATE_COMPLETE] 内容块验证完成 - 输入: {len(content_chunks)}, 有效: {len(validated_chunks)}")
        return validated_chunks

    def _safe_serialize_chunk(self, chunk):
        """安全序列化单个chunk"""
        logger.debug(f"[CHUNK_SERIALIZE] 开始序列化chunk - 类型: {type(chunk)}")
        
        try:
            if isinstance(chunk, str):
                logger.debug(f"[CHUNK_SERIALIZE_STRING] 处理字符串类型chunk")
                # 如果是字符串，尝试解析为JSON
                try:
                    parsed = json.loads(chunk)
                    result = self._safe_json_serialize(parsed)
                    logger.debug(f"[CHUNK_SERIALIZE_JSON_PARSED] JSON字符串解析并序列化成功")
                    return result
                except json.JSONDecodeError:
                    # 如果不是JSON，直接返回字符串
                    logger.debug(f"[CHUNK_SERIALIZE_PLAIN_STRING] 作为普通字符串处理")
                    return chunk
            elif isinstance(chunk, (dict, list)):
                logger.debug(f"[CHUNK_SERIALIZE_OBJECT] 处理对象类型chunk")
                # 如果是字典或列表，使用安全序列化
                result = self._safe_json_serialize(chunk)
                logger.debug(f"[CHUNK_SERIALIZE_OBJECT_SUCCESS] 对象序列化成功")
                return result
            else:
                logger.debug(f"[CHUNK_SERIALIZE_OTHER] 处理其他类型chunk，转换为字符串")
                # 其他类型转换为字符串
                return str(chunk)
        except Exception as e:
            logger.error(f"[CHUNK_SERIALIZE_ERROR] 序列化chunk失败: {str(e)}")
            # 返回错误信息
            error_content = json.dumps({"type": "markdown", "content": f"⚠️ 序列化失败: {str(e)}"})
            logger.debug(f"[CHUNK_SERIALIZE_FALLBACK] 返回错误占位符")
            return error_content

    def _calculate_content_hash(self, content):
        """计算内容哈希用于校验，增强异常处理"""
        import hashlib
        logger.debug(f"[HASH_CALCULATE] 开始计算内容哈希 - 内容类型: {type(content)}")
        
        try:
            if isinstance(content, (list, dict)):
                logger.debug(f"[HASH_SERIALIZE] 序列化复杂对象用于哈希计算")
                # 对于复杂对象，先序列化
                content_str = self._safe_json_serialize(content)
            elif isinstance(content, str):
                logger.debug(f"[HASH_STRING] 直接使用字符串计算哈希")
                content_str = content
            else:
                logger.debug(f"[HASH_CONVERT] 转换为字符串计算哈希")
                content_str = str(content)
            
            # 确保内容是UTF-8编码的字节
            content_bytes = content_str.encode('utf-8', errors='ignore')
            hash_result = hashlib.md5(content_bytes).hexdigest()
            logger.debug(f"[HASH_SUCCESS] 哈希计算成功 - 结果: {hash_result[:8]}...")
            return hash_result
            
        except Exception as e:
            logger.error(f"[HASH_ERROR] 计算哈希失败: {str(e)}")
            # 返回一个基于时间的默认哈希
            import time
            fallback_content = f"hash_error_{time.time()}_{str(content)[:50]}"
            fallback_hash = hashlib.md5(fallback_content.encode('utf-8', errors='ignore')).hexdigest()
            logger.debug(f"[HASH_FALLBACK] 使用fallback哈希 - 结果: {fallback_hash[:8]}...")
            return fallback_hash

    async def _save_ai_response_with_retry(self, session_id, response, max_retries=2):
        """保存AI回答到数据库，带重试机制"""
        logger.info(f"[AI_SAVE_RETRY] 开始保存AI响应 - 会话ID: {session_id}, 最大重试次数: {max_retries}")
        
        # 记录响应基本信息
        answer_length = len(response.get("answer", "")) if response else 0
        sources_count = len(response.get("sources", [])) if response else 0
        reply_count = len(response.get("reply", [])) if response else 0
        confidence = response.get("confidence", 0.0) if response else 0.0
        
        logger.info(f"[AI_SAVE_RETRY_INFO] 待保存响应信息 - 答案长度: {answer_length}, 来源数: {sources_count}, 富文本块: {reply_count}, 置信度: {confidence}")
        
        for attempt in range(max_retries + 1):
            try:
                logger.info(f"[AI_SAVE_ATTEMPT] 第 {attempt + 1} 次保存尝试 - 会话ID: {session_id}")
                logger.debug(f"[AI_SAVE_ATTEMPT_DETAILS] 尝试 #{attempt + 1} - 答案预览: {response.get('answer', '')[:100] if response else 'None'}{'...' if response and len(response.get('answer', '')) > 100 else ''}")
                
                ai_message = await self._save_ai_response(session_id, response)
                if ai_message:
                    logger.info(f"[AI_SAVE_SUCCESS] AI响应保存成功 (第{attempt + 1}次尝试) - 消息ID: {ai_message.id}")
                    logger.debug(f"[AI_SAVE_SUCCESS_DETAILS] 保存成功详情 - 角色: {ai_message.role}, 内容长度: {len(ai_message.content) if ai_message.content else 0}, 有分片: {ai_message.has_fragments if hasattr(ai_message, 'has_fragments') else False}")
                    return ai_message
                else:
                    logger.warning(f"[AI_SAVE_NULL] AI响应保存返回None (第{attempt + 1}次尝试)")
                    if attempt < max_retries:
                        sleep_time = 0.5 * (attempt + 1)
                        logger.info(f"[AI_SAVE_WAIT] 等待 {sleep_time} 秒后重试")
                        await asyncio.sleep(sleep_time)  # 指数退避
                        continue
                    else:
                        logger.error(f"[AI_SAVE_EXHAUSTED] 所有保存尝试均失败")
                        return None
            except Exception as e:
                logger.error(f"[AI_SAVE_ERROR] AI响应保存失败 (第{attempt + 1}次尝试): {str(e)}")
                logger.debug(f"[AI_SAVE_ERROR_CONTEXT] 保存失败上下文 - 会话ID: {session_id}, 响应类型: {type(response)}, 错误类型: {type(e).__name__}")
                
                if attempt < max_retries:
                    sleep_time = 0.5 * (attempt + 1)
                    logger.info(f"[AI_SAVE_RETRY_WAIT] 等待 {sleep_time} 秒后重试")
                    await asyncio.sleep(sleep_time)  # 指数退避
                    continue
                else:
                    # 最后一次尝试失败，记录到Redis或日志系统
                    logger.error(f"[AI_SAVE_FINAL_FAIL] 所有保存尝试均失败，记录失败日志")
                    logger.debug(f"[AI_SAVE_FINAL_CONTEXT] 最终失败上下文 - 总尝试次数: {max_retries + 1}, 最后错误: {str(e)}")
                    await self._log_failed_save(session_id, response, str(e))
                    return None
        
        logger.error(f"[AI_SAVE_RETRY_COMPLETE] AI响应保存重试流程结束 - 未成功保存")
        return None

    async def _log_failed_save(self, session_id, response, error_msg):
        """记录保存失败的消息，用于后续补偿"""
        logger.info(f"[FAILED_SAVE_LOG] 开始记录保存失败的消息 - 会话ID: {session_id}")
        
        try:
            # 这里可以集成Redis或其他队列系统
            failed_save_data = {
                "session_id": session_id,
                "response": response,
                "error": error_msg,
                "timestamp": datetime.utcnow().isoformat(),
                "retry_count": 0
            }
            
            logger.debug(f"[FAILED_SAVE_DATA] 构建失败保存数据 - 错误: {error_msg[:100]}...")
            
            # 简单的文件日志记录（生产环境应该使用Redis）
            import json
            log_file = "failed_saves.log"
            
            logger.debug(f"[FAILED_SAVE_FILE] 写入失败日志文件: {log_file}")
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(json.dumps(failed_save_data, ensure_ascii=False) + "\n")
            
            logger.info(f"[FAILED_SAVE_LOGGED] 保存失败记录已写入日志 - 会话ID: {session_id}")
        except Exception as log_error:
            logger.error(f"[FAILED_SAVE_LOG_ERROR] 记录保存失败日志时出错: {str(log_error)}")
            # 即使日志记录失败，也不应该影响主流程

    async def _save_ai_response(self, session_id, response):
        """保存AI回答到数据库，增强大型内容处理和序列化安全性"""
        try:
            logger.info(f"[AI_SAVE_START] 开始保存AI响应 - 会话ID: {session_id}")
            
            # 记录响应内容的详细信息
            answer = response.get("answer", "")
            sources = response.get("sources", [])
            confidence = response.get("confidence", 0.0)
            reply = response.get("reply", [])
            
            logger.info(f"[AI_RESPONSE_DETAILS] AI响应详细信息 - 答案长度: {len(answer)} 字符, 来源数量: {len(sources)}, 置信度: {confidence}, 富文本块数: {len(reply)}")
            logger.debug(f"[AI_RESPONSE_PREVIEW] 答案内容预览: {answer[:200]}{'...' if len(answer) > 200 else ''}")
            
            if sources:
                logger.debug(f"[AI_RESPONSE_SOURCES] 来源信息: {[s.get('line_number', 'N/A') for s in sources[:3]]}{'...' if len(sources) > 3 else ''}")
            
            if reply and len(reply) > 0:
                logger.debug(f"[AI_RESPONSE_RICH_TYPES] 富文本类型分布: {[item.get('type', 'unknown') for item in reply[:5]]}")
                
                # 安全序列化，避免循环引用
                try:
                    serialized = self._safe_json_serialize(reply)
                    content_size = len(serialized)
                    logger.info(f"[AI_SAVE_SERIALIZE] 富文本内容序列化成功 - 大小: {content_size} bytes")
                    
                    # 记录保存策略选择
                    if content_size > 100000:  # 约100KB
                        logger.warning(f"[AI_SAVE_STRATEGY] 选择分片存储策略 - 内容大小: {content_size} bytes 超过阈值")
                    else:
                        logger.info(f"[AI_SAVE_STRATEGY] 选择常规存储策略 - 内容大小: {content_size} bytes")
                        
                except Exception as serialize_error:
                    logger.error(f"[AI_SAVE_SERIALIZE_ERROR] 序列化失败，使用fallback: {str(serialize_error)}")
                    # Fallback: 转换为安全格式
                    reply = self._convert_to_safe_format(reply)
                    serialized = self._safe_json_serialize(reply)
                    content_size = len(serialized)
                    logger.info(f"[AI_SAVE_FALLBACK_SERIALIZE] Fallback序列化成功 - 大小: {content_size} bytes")
                
                # 如果超过阈值，使用独立事务拆分存储
                if content_size > 100000:  # 约100KB
                    logger.warning(f"[AI_SAVE_LARGE] 内容超过大小限制，使用分片存储 - 大小: {content_size} bytes")
                    
                    # 记录分片存储过程
                    logger.info(f"[AI_SAVE_FRAGMENT_START] 开始分片存储过程")
                    
                    # 创建主消息，包含摘要
                    main_message = ChatMessage(
                        session_id=uuid.UUID(session_id),
                        role='assistant',
                        content=response.get("answer", "")[:5000],
                        sources=self._sanitize_sources(response.get("sources", [])[:5]),
                        confidence=float(response.get("confidence", 0.0)),
                        message_type='markdown',
                        rich_content=[{
                            "type": "markdown", 
                            "content": "此回复包含大量数据，已分片存储以提高性能"
                        }],
                        has_fragments=True,
                        created_at=datetime.utcnow()
                    )
                    self.db.add(main_message)
                    
                    # 立即提交主消息
                    self.db.commit()
                    logger.info(f"[AI_SAVE_MAIN] 主消息创建并提交成功 - 消息ID: {main_message.id}")
                    
                    # 使用独立事务处理分片存储
                    logger.info(f"[AI_SAVE_FRAGMENTS] 开始分片存储")
                    chunks = self.split_content_safely(reply)
                    logger.info(f"[AI_SAVE_CHUNKS] 内容分片完成 - 分片数量: {len(chunks)}")
                    
                    fragments_stored = await self._save_message_fragments(main_message.id, chunks)
                    
                    if not fragments_stored:
                        # 如果片段存储失败，添加警告到主消息
                        logger.error(f"[AI_SAVE_FRAGMENTS_FAILED] 分片存储失败，添加警告")
                        main_message.rich_content.append({
                            "type": "markdown",
                            "content": "⚠️ 警告：消息片段存储失败，内容可能不完整"
                        })
                        self.db.commit()  # 提交警告更新
                        logger.warning(f"[AI_SAVE_WARNING_ADDED] 已添加分片存储失败警告")
                    else:
                        logger.info(f"[AI_SAVE_FRAGMENTS_SUCCESS] 分片存储成功")
                    
                    logger.info(f"[AI_SAVE_FRAGMENT_COMPLETE] 分片存储过程完成 - 主消息ID: {main_message.id}")
                    return main_message
            
            # 常规消息处理逻辑
            logger.info(f"[AI_SAVE_REGULAR] 使用常规消息存储")
            logger.debug(f"[AI_SAVE_REGULAR_DETAILS] 常规存储 - 内容长度: {len(answer)}, 来源数: {len(sources)}, 富文本块: {len(reply)}")
            
            ai_time = datetime.utcnow()
            assistant_message = ChatMessage(
                session_id=uuid.UUID(session_id),
                role='assistant',
                content=response.get("answer", ""),
                sources=self._sanitize_sources(response.get("sources", [])),
                confidence=float(response.get("confidence", 0.0)),
                message_type='markdown',
                rich_content=reply,
                created_at=ai_time
            )
            self.db.add(assistant_message)
            
            # 立即提交AI消息
            self.db.commit()
            logger.info(f"[AI_SAVE_REGULAR_SUCCESS] 常规消息创建并提交成功 - 消息ID: {assistant_message.id}")
            logger.debug(f"[AI_SAVE_REGULAR_TIMESTAMP] 消息时间戳: {ai_time}")
            
            return assistant_message
        except Exception as e:
            self.db.rollback()
            logger.error(f"[AI_SAVE_EXCEPTION] 保存AI响应时发生异常: {str(e)}")
            import traceback
            logger.debug(f"[AI_SAVE_TRACEBACK] 保存异常堆栈: {traceback.format_exc()}")
            return None

    def _safe_json_serialize(self, data):
        """安全的JSON序列化，处理循环引用和特殊类型"""
        logger.debug(f"[JSON_SERIALIZE] 开始安全JSON序列化 - 数据类型: {type(data)}")
        
        try:
            # 使用自定义编码器处理特殊类型
            result = json.dumps(data, ensure_ascii=False, default=self._json_serializer, separators=(',', ':'))
            logger.debug(f"[JSON_SERIALIZE_SUCCESS] JSON序列化成功 - 结果长度: {len(result)} 字符")
            return result
        except (TypeError, ValueError) as e:
            logger.warning(f"[JSON_SERIALIZE_ERROR] JSON序列化失败，进行深度清理: {str(e)}")
            # 深度清理数据
            cleaned_data = self._deep_clean_for_json(data)
            result = json.dumps(cleaned_data, ensure_ascii=False, separators=(',', ':'))
            logger.debug(f"[JSON_SERIALIZE_CLEANED] 清理后序列化成功 - 结果长度: {len(result)} 字符")
            return result

    def _json_serializer(self, obj):
        """自定义JSON序列化器，处理特殊类型"""
        logger.debug(f"[JSON_CUSTOM_SERIALIZER] 处理特殊类型对象 - 类型: {type(obj)}")
        
        if hasattr(obj, 'isoformat'):  # datetime对象
            result = obj.isoformat()
            logger.debug(f"[JSON_DATETIME] 序列化datetime对象: {result}")
            return result
        elif hasattr(obj, '__dict__'):  # 自定义对象
            result = str(obj)
            logger.debug(f"[JSON_CUSTOM_OBJECT] 序列化自定义对象为字符串")
            return result
        elif isinstance(obj, (set, frozenset)):
            result = list(obj)
            logger.debug(f"[JSON_SET] 序列化set/frozenset为列表 - 长度: {len(result)}")
            return result
        elif isinstance(obj, bytes):
            result = obj.decode('utf-8', errors='ignore')
            logger.debug(f"[JSON_BYTES] 序列化bytes对象 - 长度: {len(result)} 字符")
            return result
        else:
            result = str(obj)
            logger.debug(f"[JSON_OTHER] 序列化其他类型为字符串")
            return result

    def _deep_clean_for_json(self, data, max_depth=10, current_depth=0):
        """深度清理数据，移除循环引用和不可序列化的对象"""
        logger.debug(f"[JSON_DEEP_CLEAN] 深度清理数据 - 当前深度: {current_depth}, 最大深度: {max_depth}, 数据类型: {type(data)}")
        
        if current_depth > max_depth:
            logger.warning(f"[JSON_MAX_DEPTH] 达到最大深度限制: {max_depth}")
            return "... (max depth reached)"
        
        if isinstance(data, dict):
            logger.debug(f"[JSON_CLEAN_DICT] 清理字典 - 键数量: {len(data)}")
            cleaned = {}
            for key, value in data.items():
                try:
                    # 确保key是字符串
                    safe_key = str(key) if not isinstance(key, str) else key
                    cleaned[safe_key] = self._deep_clean_for_json(value, max_depth, current_depth + 1)
                except Exception as e:
                    logger.error(f"[JSON_CLEAN_DICT_ITEM] 清理字典项失败 {key}: {str(e)}")
                    cleaned[str(key)] = str(value)
            logger.debug(f"[JSON_CLEAN_DICT_DONE] 字典清理完成 - 清理后键数量: {len(cleaned)}")
            return cleaned
        elif isinstance(data, (list, tuple)):
            logger.debug(f"[JSON_CLEAN_LIST] 清理列表/元组 - 项目数量: {len(data)}")
            cleaned = []
            for i, item in enumerate(data):
                try:
                    cleaned.append(self._deep_clean_for_json(item, max_depth, current_depth + 1))
                except Exception as e:
                    logger.error(f"[JSON_CLEAN_LIST_ITEM] 清理列表项 #{i} 失败: {str(e)}")
                    cleaned.append(str(item))
            logger.debug(f"[JSON_CLEAN_LIST_DONE] 列表清理完成 - 清理后项目数量: {len(cleaned)}")
            return cleaned
        elif isinstance(data, (str, int, float, bool, type(None))):
            logger.debug(f"[JSON_CLEAN_PRIMITIVE] 基础类型无需清理")
            return data
        else:
            # 对于其他类型，转换为字符串
            result = str(data)
            logger.debug(f"[JSON_CLEAN_OTHER] 其他类型转换为字符串 - 长度: {len(result)} 字符")
            return result

    def _convert_to_safe_format(self, reply):
        """将回复转换为安全的格式"""
        logger.info(f"[SAFE_FORMAT] 开始转换回复为安全格式 - 输入类型: {type(reply)}")
        
        if not isinstance(reply, list):
            logger.warning(f"[SAFE_FORMAT_NOT_LIST] 输入不是列表，转换为单项列表")
            return [{"type": "markdown", "content": str(reply)}]
        
        logger.debug(f"[SAFE_FORMAT_LIST] 处理列表回复 - 项目数量: {len(reply)}")
        safe_reply = []
        for i, item in enumerate(reply):
            try:
                logger.debug(f"[SAFE_FORMAT_ITEM] 处理回复项 #{i} - 类型: {type(item)}")
                
                if isinstance(item, dict):
                    safe_item = {
                        "type": str(item.get("type", "markdown")),
                        "content": str(item.get("content", ""))
                    }
                    
                    # 安全处理可选字段
                    if "language" in item:
                        safe_item["language"] = str(item["language"])[:50]  # 限制长度
                        logger.debug(f"[SAFE_FORMAT_LANGUAGE] 添加语言字段: {safe_item['language']}")
                    if "columns" in item and isinstance(item["columns"], list):
                        safe_item["columns"] = [str(col)[:100] for col in item["columns"][:20]]  # 限制数量和长度
                        logger.debug(f"[SAFE_FORMAT_COLUMNS] 添加列字段 - 数量: {len(safe_item['columns'])}")
                    if "rows" in item and isinstance(item["rows"], list):
                        safe_rows = []
                        for row_idx, row in enumerate(item["rows"][:100]):  # 最多100行
                            if isinstance(row, list):
                                safe_row = [str(cell)[:200] for cell in row[:20]]  # 最多20列，每个单元格最长200字符
                                safe_rows.append(safe_row)
                        safe_item["rows"] = safe_rows
                        logger.debug(f"[SAFE_FORMAT_ROWS] 添加行数据 - 行数: {len(safe_rows)}")
                    if "metadata" in item and isinstance(item["metadata"], dict):
                        safe_item["metadata"] = self._deep_clean_for_json(item["metadata"], max_depth=3)
                        logger.debug(f"[SAFE_FORMAT_METADATA] 添加元数据字段")
                    
                    safe_reply.append(safe_item)
                    logger.debug(f"[SAFE_FORMAT_ITEM_SUCCESS] 回复项 #{i} 处理成功")
                else:
                    safe_reply.append({"type": "markdown", "content": str(item)})
                    logger.debug(f"[SAFE_FORMAT_ITEM_FALLBACK] 回复项 #{i} 使用fallback格式")
            except Exception as e:
                logger.error(f"[SAFE_FORMAT_ITEM_ERROR] 转换回复项 #{i} 失败: {str(e)}")
                safe_reply.append({"type": "markdown", "content": str(item)})
        
        logger.info(f"[SAFE_FORMAT_COMPLETE] 安全格式转换完成 - 输入: {len(reply)}, 输出: {len(safe_reply)}")
        return safe_reply

    def _sanitize_sources(self, sources):
        """清理sources数据，确保可序列化"""
        logger.info(f"[SOURCES_SANITIZE] 开始清理sources数据 - 输入类型: {type(sources)}")
        
        if not isinstance(sources, list):
            logger.warning(f"[SOURCES_NOT_LIST] sources不是列表，返回空列表")
            return []
        
        logger.debug(f"[SOURCES_PROCESS] 处理sources列表 - 原始数量: {len(sources)}")
        sanitized = []
        for i, source in enumerate(sources[:10]):  # 最多10个来源
            try:
                logger.debug(f"[SOURCES_ITEM] 处理source #{i} - 类型: {type(source)}")
                
                if isinstance(source, dict):
                    sanitized_source = {
                        "line_number": int(source.get("line_number", 0)),
                        "content": str(source.get("content", ""))[:500],  # 限制内容长度
                        "page": int(source.get("page", 1)),
                        "start_pos": int(source.get("start_pos", 0)),
                        "end_pos": int(source.get("end_pos", 0)),
                        "is_scanned": bool(source.get("is_scanned", False)),
                        "similarity": float(source.get("similarity", 0.0))
                    }
                    
                    # 可选字段
                    if "document_id" in source:
                        sanitized_source["document_id"] = str(source["document_id"])
                        logger.debug(f"[SOURCES_DOC_ID] 添加文档ID: {sanitized_source['document_id']}")
                    if "document_name" in source:
                        sanitized_source["document_name"] = str(source["document_name"])[:100]
                        logger.debug(f"[SOURCES_DOC_NAME] 添加文档名称: {sanitized_source['document_name'][:20]}...")
                    
                    sanitized.append(sanitized_source)
                    logger.debug(f"[SOURCES_ITEM_SUCCESS] source #{i} 处理成功")
            except Exception as e:
                logger.error(f"[SOURCES_ITEM_ERROR] 清理source #{i} 失败: {str(e)}")
                continue
        
        logger.info(f"[SOURCES_SANITIZE_COMPLETE] sources清理完成 - 输入: {len(sources)}, 输出: {len(sanitized)}")
        return sanitized

    def split_content_safely(self, content, max_size=50000):
        """安全地将大型内容拆分成多个片段，避免序列化问题"""
        logger.info(f"[CONTENT_SPLIT] 开始安全拆分内容 - 内容类型: {type(content)}, 最大大小: {max_size} bytes")
        
        if isinstance(content, list):
            logger.debug(f"[CONTENT_SPLIT_LIST] 处理列表内容 - 项目数量: {len(content)}")
            result = []
            current_chunk = []
            current_size = 0
            
            for i, item in enumerate(content):
                try:
                    logger.debug(f"[CONTENT_SPLIT_ITEM] 处理内容项 #{i}")
                    # 先清理item，确保可序列化
                    safe_item = self._deep_clean_for_json(item, max_depth=5)
                    item_serialized = self._safe_json_serialize(safe_item)
                    item_size = len(item_serialized)
                    
                    logger.debug(f"[CONTENT_SPLIT_SIZE] 内容项 #{i} 大小: {item_size} bytes")
                
                    if current_size + item_size > max_size and current_chunk:
                        # 当前块已满，保存并创建新块
                        logger.debug(f"[CONTENT_SPLIT_CHUNK_FULL] 当前块已满，创建新块 - 当前大小: {current_size}, 项目数: {len(current_chunk)}")
                        result.append(current_chunk)
                        current_chunk = [safe_item]
                        current_size = item_size
                    else:
                        # 添加到当前块
                        current_chunk.append(safe_item)
                        current_size += item_size
                        logger.debug(f"[CONTENT_SPLIT_ADD] 添加到当前块 - 新大小: {current_size}")
                        
                except Exception as e:
                    logger.error(f"[CONTENT_SPLIT_ITEM_ERROR] 处理内容项 #{i} 失败: {str(e)}")
                    # 添加错误占位符
                    error_item = {"type": "markdown", "content": f"⚠️ 内容处理失败: {str(e)}"}
                    current_chunk.append(error_item)
                    current_size += 100  # 估算大小
            
            # 添加最后一个块
            if current_chunk:
                logger.debug(f"[CONTENT_SPLIT_FINAL_CHUNK] 添加最后一个块 - 大小: {current_size}, 项目数: {len(current_chunk)}")
                result.append(current_chunk)
            
            logger.info(f"[CONTENT_SPLIT_LIST_COMPLETE] 列表内容拆分完成 - 生成块数: {len(result)}")
            return result
        else:
            logger.debug(f"[CONTENT_SPLIT_OTHER] 处理非列表内容，转换为安全格式")
            # 如果不是列表，转换为安全格式
            safe_content = self._deep_clean_for_json(content, max_depth=5)
            logger.info(f"[CONTENT_SPLIT_OTHER_COMPLETE] 非列表内容处理完成")
            return [safe_content]

    def _format_message(self, message):
        """将消息对象格式化为API响应格式"""
        logger.debug(f"[FORMAT_MESSAGE] 开始格式化用户消息 - 消息对象: {message is not None}")
        
        if not message:
            # 生成更具唯一性的ID
            unique_id = f"msg_user_{datetime.utcnow().isoformat()}_{uuid.uuid4().hex[:8]}"
            logger.warning(f"[FORMAT_MESSAGE_EMPTY] 消息对象为空，生成默认格式 - ID: {unique_id}")
            return {
                "id": unique_id,
                "role": "user",
                "content": "",
                "created_at": datetime.utcnow().isoformat(),
                "sources": [],
                "confidence": 0,
                "reply": []
            }
        
        logger.debug(f"[FORMAT_MESSAGE_SUCCESS] 用户消息格式化成功 - 消息ID: {message.id}, 角色: {message.role}, 内容长度: {len(message.content) if message.content else 0}")
        
        return {
            "id": str(message.id),
            "role": message.role,
            "content": message.content,
            "created_at": message.created_at.isoformat() if message.created_at else datetime.utcnow().isoformat(),
            "sources": [],
            "confidence": 0,
            "reply": []
        }

    def _format_ai_message(self, response, message=None):
        """将AI响应格式化为API响应格式"""
        logger.debug(f"[FORMAT_AI_MESSAGE] 开始格式化AI消息 - 有响应: {response is not None}, 有消息对象: {message is not None}")
        
        if message:
            logger.debug(f"[FORMAT_AI_MESSAGE_FROM_DB] 从数据库消息对象格式化 - 消息ID: {message.id}, 内容长度: {len(message.content) if message.content else 0}")
            logger.debug(f"[FORMAT_AI_MESSAGE_SOURCES] 消息来源数量: {len(message.sources) if message.sources else 0}")
            logger.debug(f"[FORMAT_AI_MESSAGE_RICH] 富文本内容块数: {len(message.rich_content) if message.rich_content else 0}")
            
            return {
                "id": str(message.id),
                "role": "assistant",
                "content": message.content,
                "created_at": message.created_at.isoformat() if message.created_at else datetime.utcnow().isoformat(),
                "sources": message.sources or [],
                "confidence": message.confidence or 0,
                "reply": message.rich_content or [{"type": "markdown", "content": message.content or ""}]
            }
        
        # 生成更具唯一性的ID
        unique_id = f"msg_ai_{datetime.utcnow().isoformat()}_{uuid.uuid4().hex[:8]}"
        logger.debug(f"[FORMAT_AI_MESSAGE_GENERATED] 生成临时AI消息格式 - ID: {unique_id}")
        
        # 修复：处理response为None的情况
        if response is None:
            logger.warning(f"[FORMAT_AI_MESSAGE_NULL] 响应对象为空，生成错误消息格式")
            return {
                "id": unique_id,
                "role": "assistant",
                "content": "处理请求时发生错误",
                "created_at": datetime.utcnow().isoformat(),
                "sources": [],
                "confidence": 0.0,
                "reply": [{"type": "markdown", "content": "处理请求时发生错误"}]
            }
        
        # 使用响应对象创建格式化消息
        answer = response.get("answer", "")
        sources = response.get("sources", [])
        confidence = response.get("confidence", 0.0)
        reply = response.get("reply", [{"type": "markdown", "content": answer}])
        
        logger.debug(f"[FORMAT_AI_MESSAGE_FROM_RESPONSE] 从响应对象格式化 - 答案长度: {len(answer)}, 来源数: {len(sources)}, 置信度: {confidence}")
        logger.debug(f"[FORMAT_AI_MESSAGE_REPLY] 富文本回复块数: {len(reply)}")
        
        return {
            "id": unique_id,
            "role": "assistant",
            "content": answer,
            "created_at": datetime.utcnow().isoformat(),
            "sources": sources,
            "confidence": confidence,
            "reply": reply
        }

    async def create_ai_chat_session(self, user_id: str, title: str = None) -> dict:
        """创建一个无文档的普通AI对话会话"""
        logger.info(f"[CREATE_AI_SESSION] 开始创建AI对话会话 - 用户ID: {user_id}, 标题: {title}")
        
        default_title = title or f"AI对话 {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
        logger.debug(f"[CREATE_AI_SESSION_TITLE] 使用标题: {default_title}")
        
        result = await self.create_chat_session(
            user_id=user_id,
            title=default_title,
            paper_ids=None,
            session_type="general"
        )
        
        logger.info(f"[CREATE_AI_SESSION_SUCCESS] AI对话会话创建成功 - 会话ID: {result.get('id')}")
        return result

    async def remove_document_from_session(self, session_id: str, document_id: str, user_id: str) -> dict:
        """从会话中移除文档"""
        logger.info(f"[REMOVE_DOC_FROM_SESSION] 开始从会话移除文档 - 会话ID: {session_id}, 文档ID: {document_id}, 用户ID: {user_id}")
        
        try:
            # 验证会话存在且属于该用户
            logger.debug(f"[REMOVE_DOC_VERIFY_SESSION] 验证会话存在且属于用户")
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                logger.error(f"[REMOVE_DOC_SESSION_NOT_FOUND] 会话不存在或不属于该用户 - 会话ID: {session_id}, 用户ID: {user_id}")
                raise ValueError("会话不存在或不属于该用户")
            
            logger.debug(f"[REMOVE_DOC_SESSION_FOUND] 会话验证成功 - 标题: {session.title}")
            
            # 查找要删除的文档关联
            logger.debug(f"[REMOVE_DOC_FIND_ASSOCIATION] 查找文档关联")
            doc = self.db.query(SessionDocument).filter(
                SessionDocument.id == uuid.UUID(document_id),
                SessionDocument.session_id == uuid.UUID(session_id)
            ).first()
            
            if not doc:
                logger.error(f"[REMOVE_DOC_NOT_IN_SESSION] 文档不在当前会话中或不存在 - 文档ID: {document_id}")
                raise ValueError("文档不在当前会话中或不存在")
            
            logger.debug(f"[REMOVE_DOC_ASSOCIATION_FOUND] 找到文档关联 - 文档名: {doc.filename}")
            
            # 获取文档数量
            doc_count = self.db.query(SessionDocument).filter(
                SessionDocument.session_id == uuid.UUID(session_id)
            ).count()
            
            logger.info(f"[REMOVE_DOC_COUNT] 会话当前文档数量: {doc_count}")
            
            # 如果是唯一文档，需要特殊处理
            if doc_count == 1:
                logger.info(f"[REMOVE_DOC_LAST_DOC] 移除最后一个文档，会话类型将变为通用")
                # 如果是最后一个文档，清除session.paper_id并修改类型
                session.paper_id = None
                session.session_type = "general"
            elif session.paper_id == doc.paper_id:
                logger.info(f"[REMOVE_DOC_MAIN_DOC] 移除主文档，需要更新为另一个文档")
                # 如果删除的是主文档，更新为另一个文档
                next_doc = self.db.query(SessionDocument).filter(
                    SessionDocument.session_id == uuid.UUID(session_id),
                    SessionDocument.id != uuid.UUID(document_id)
                ).first()
                
                if next_doc:
                    session.paper_id = next_doc.paper_id
                    logger.debug(f"[REMOVE_DOC_NEW_MAIN] 新主文档设置为: {next_doc.filename}")
            
            # 删除文档关联
            logger.debug(f"[REMOVE_DOC_DELETE_ASSOCIATION] 删除文档关联")
            self.db.delete(doc)
            
            # 更新其他文档的顺序
            logger.debug(f"[REMOVE_DOC_REORDER] 重新排序剩余文档")
            remaining_docs = self.db.query(SessionDocument).filter(
                SessionDocument.session_id == uuid.UUID(session_id)
            ).order_by(SessionDocument.order).all()
            
            for i, remaining_doc in enumerate(remaining_docs):
                remaining_doc.order = i
            
            logger.debug(f"[REMOVE_DOC_REORDER_COMPLETE] 文档重新排序完成 - 剩余文档数: {len(remaining_docs)}")
            
            self.db.commit()
            logger.info(f"[REMOVE_DOC_SUCCESS] 文档移除成功 - 会话类型: {session.session_type}")
            
            return {
                "status": "success",
                "message": "Document removed from session",
                "session_id": session_id,
                "session_type": session.session_type
            }
        except Exception as e:
            self.db.rollback()
            logger.error(f"[REMOVE_DOC_ERROR] 从会话移除文档失败: {str(e)}")
            raise Exception(f"Failed to remove document from session: {str(e)}")

    async def get_session_documents(self, session_id: str, user_id: str) -> List[dict]:
        """获取会话关联的所有文档"""
        try:
            # 验证会话存在且属于该用户
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                raise ValueError("会话不存在或不属于该用户")
            
            # 获取关联文档
            session_docs = self.db.query(SessionDocument).filter(
                SessionDocument.session_id == uuid.UUID(session_id)
            ).order_by(SessionDocument.order).all()
            
            # 构建结果
            result = []
            for doc in session_docs:
                # 获取文档详情
                paper = self.db.query(PaperAnalysis).filter(
                    PaperAnalysis.paper_id == doc.paper_id
                ).first()
                
                if paper:
                    result.append({
                        "id": str(doc.id),
                        "paper_id": str(doc.paper_id),
                        "filename": paper.filename,
                        "file_type": paper.file_type,
                        "file_size": paper.file_size,
                        "order": doc.order,
                        "added_at": doc.added_at.isoformat() if doc.added_at else None,
                        "summary": paper.summary
                    })
            
            return result
        except Exception as e:
            raise Exception(f"Failed to get session documents: {str(e)}")

    async def add_document_to_session(self, session_id: str, paper_id: str, user_id: str) -> dict:
        """向现有会话添加文档"""
        try:
            # 验证会话存在且属于该用户
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                raise ValueError("会话不存在或不属于该用户")
            
            # 验证文档存在
            paper = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.paper_id == uuid.UUID(paper_id)
            ).first()
            
            if not paper:
                raise ValueError(f"文档ID {paper_id} 不存在")
            
            # 检查该文档是否已在会话中
            existing_doc = self.db.query(SessionDocument).filter(
                SessionDocument.session_id == uuid.UUID(session_id),
                SessionDocument.paper_id == uuid.UUID(paper_id)
            ).first()
            
            if existing_doc:
                raise ValueError("该文档已添加到会话中")
            
            # 检查会话中的文档数量是否超出限制
            doc_count = self.db.query(SessionDocument).filter(
                SessionDocument.session_id == uuid.UUID(session_id)
            ).count()
            
            if doc_count >= self.MAX_DOCUMENTS_PER_SESSION:
                raise ValueError(f"一个会话最多支持{self.MAX_DOCUMENTS_PER_SESSION}个文档")
            
            # 创建关联
            session_doc = SessionDocument(
                session_id=uuid.UUID(session_id),
                paper_id=uuid.UUID(paper_id),
                order=doc_count,  # 新文档添加到末尾
                filename=paper.filename
            )
            self.db.add(session_doc)
            
            # 如果是首个文档，更新会话的paper_id（向后兼容）
            if doc_count == 0 and not session.paper_id:
                session.paper_id = uuid.UUID(paper_id)
                session.session_type = "document"  # 更新类型为文档会话
            
            self.db.commit()
            
            # 返回添加的文档信息
            return {
                "id": str(session_doc.id),
                "session_id": session_id,
                "paper_id": paper_id,
                "filename": paper.filename,
                "order": session_doc.order
            }
        except Exception as e:
            self.db.rollback()
            raise Exception(f"Failed to add document to session: {str(e)}")

    async def delete_chat_session(self, session_id: str, user_id: str) -> dict:
        """逻辑删除会话（仅标记为非活跃）"""
        try:
            # 验证会话存在且属于该用户
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                raise ValueError("会话不存在或不属于该用户")
            
            # 执行逻辑删除（标记为非活跃）
            session.is_active = False
            session.updated_at = datetime.utcnow()
            
            self.db.commit()
            
            return {
                "status": "success",
                "message": "会话已删除",
                "session_id": session_id
            }
        except Exception as e:
            self.db.rollback()
            raise Exception(f"删除会话失败: {str(e)}")

    async def update_chat_session_title(self, session_id: str, user_id: str, new_title: str) -> dict:
        """更新会话标题"""
        logger.info(f"[UPDATE_SESSION_TITLE] 开始更新会话标题 - 会话ID: {session_id}, 用户ID: {user_id}, 新标题: {new_title}")
        
        try:
            # 验证会话存在且属于该用户
            logger.debug(f"[UPDATE_TITLE_VERIFY] 验证会话存在且属于用户")
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                logger.error(f"[UPDATE_TITLE_NOT_FOUND] 会话不存在或不属于该用户 - 会话ID: {session_id}, 用户ID: {user_id}")
                raise ValueError("会话不存在或不属于该用户")
            
            old_title = session.title
            logger.debug(f"[UPDATE_TITLE_OLD] 原标题: {old_title}")
            
            # 更新标题
            session.title = new_title
            session.updated_at = datetime.utcnow()
            
            logger.debug(f"[UPDATE_TITLE_SET] 设置新标题和更新时间")
            
            self.db.commit()
            logger.info(f"[UPDATE_TITLE_SUCCESS] 会话标题更新成功 - 从 '{old_title}' 更新为 '{new_title}'")
            
            return {
                "id": str(session.id),
                "title": session.title,
                "updated_at": session.updated_at.isoformat(),
                "status": "success"
            }
        except Exception as e:
            self.db.rollback()
            logger.error(f"[UPDATE_TITLE_ERROR] 更新会话标题失败: {str(e)}")
            raise Exception(f"更新会话标题失败: {str(e)}")

    def _estimate_memory_usage(self, content_size):
        """预估内存使用并设置处理级别"""
        logger.debug(f"[MEMORY_ESTIMATE] 开始预估内存使用 - 内容大小: {content_size} bytes")
        
        # 预估内存使用并设置处理级别
        estimated_mb = content_size * 5 / (1024 * 1024)  # 粗略估计
        
        if estimated_mb > 500:
            level = "high"
            logger.warning(f"[MEMORY_ESTIMATE_HIGH] 高内存使用预警 - 预估: {estimated_mb:.2f} MB")
        elif estimated_mb > 100:
            level = "medium"
            logger.info(f"[MEMORY_ESTIMATE_MEDIUM] 中等内存使用 - 预估: {estimated_mb:.2f} MB")
        else:
            level = "low"
            logger.debug(f"[MEMORY_ESTIMATE_LOW] 低内存使用 - 预估: {estimated_mb:.2f} MB")
        
        result = {
            "level": level,
            "estimated_mb": estimated_mb
        }
        
        logger.debug(f"[MEMORY_ESTIMATE_COMPLETE] 内存预估完成 - 级别: {level}, 大小: {estimated_mb:.2f} MB")
        return result

    @contextmanager
    def reliable_transaction(self):
        """提供事务重试与恢复机制"""
        logger.debug(f"[TRANSACTION_START] 开始可靠事务处理")
        retry_count = 0
        
        while retry_count < 3:
            try:
                logger.debug(f"[TRANSACTION_ATTEMPT] 事务尝试 #{retry_count + 1}")
                yield
                self.db.commit()
                logger.info(f"[TRANSACTION_SUCCESS] 事务提交成功 - 尝试次数: {retry_count + 1}")
                return
            except Exception as e:
                self.db.rollback()
                retry_count += 1
                logger.warning(f"[TRANSACTION_RETRY] 事务失败，准备重试 - 尝试: {retry_count}/3, 错误: {str(e)}")
                
                if retry_count == 3:
                    logger.error(f"[TRANSACTION_EXHAUSTED] 事务重试次数耗尽 - 最终失败: {str(e)}")
                    raise
                
                logger.debug(f"[TRANSACTION_WAIT] 等待1秒后重试")
                time.sleep(1)  # 指数退避

    async def stream_message(self, session_id: str, message: str, user_id: str):
        """流式发送消息，逐步返回AI回复 - 使用渐进式保存机制"""
        # 添加消息状态跟踪
        message_status = {"stored": False, "user_message_id": None, "ai_message_id": None}
        ai_message = None  # 用于存储AI消息对象
        
        logger.info(f"[STREAM_START] 开始流式消息处理 - 会话ID: {session_id}, 用户ID: {user_id}, 消息长度: {len(message)} 字符")
        logger.debug(f"[STREAM_CONTENT] 流式消息内容预览: {message[:100]}{'...' if len(message) > 100 else ''}")
        
        try:
            # 1. 保存用户消息
            logger.info(f"[STREAM_USER_SAVE] 开始保存用户消息 - 会话ID: {session_id}")
            current_time = datetime.utcnow()
            user_message = ChatMessage(
                session_id=uuid.UUID(session_id),
                role='user',
                content=message,
                message_type='markdown',
                created_at=current_time
            )
            self.db.add(user_message)
            self.db.commit()  # 立即提交用户消息
            message_status["user_message_id"] = str(user_message.id)
            message_status["stored"] = True
            logger.info(f"[STREAM_USER_SAVED] 用户消息保存成功 - 消息ID: {user_message.id}")
            
            # 2. 获取会话信息
            logger.info(f"[STREAM_SESSION] 开始获取会话信息 - 会话ID: {session_id}")
            session = self.db.query(ChatSession).filter(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.user_id == uuid.UUID(user_id)
            ).first()
            
            if not session:
                logger.error(f"[STREAM_SESSION_ERROR] 会话不存在 - 会话ID: {session_id}")
                yield {"error": "Chat session not found", "done": True}
                return
            
            logger.info(f"[STREAM_SESSION_FOUND] 会话信息获取成功 - 标题: {session.title}, 类型: {session.session_type}")
            
            # 3. 预创建AI消息记录 - 渐进式保存的关键
            logger.info(f"[STREAM_AI_PRECREATE] 预创建AI消息记录")
            try:
                ai_message = ChatMessage(
                    session_id=uuid.UUID(session_id),
                    role='assistant',
                    content="[STREAMING...]",  # 使用特殊标记表示流式状态
                    message_type='streaming',  # 使用message_type标记流式状态
                    created_at=datetime.utcnow(),
                    sources=[],
                    confidence=0.0,
                    rich_content=[]
                )
                self.db.add(ai_message)
                self.db.commit()  # 立即提交，确保记录存在
                message_status["ai_message_id"] = str(ai_message.id)
                logger.info(f"[STREAM_AI_PRECREATED] AI消息预创建成功 - 消息ID: {ai_message.id}")
            except Exception as precreate_error:
                logger.error(f"[STREAM_AI_PRECREATE_ERROR] AI消息预创建失败: {str(precreate_error)}")
                # 继续处理，但没有渐进式保存
            
            # 4. 获取历史消息作为上下文
            try:
                logger.info(f"[STREAM_HISTORY] 开始获取历史消息 - 会话ID: {session_id}")
                history_messages = self.db.query(ChatMessage).filter(
                    ChatMessage.session_id == uuid.UUID(session_id),
                    ChatMessage.message_type != 'streaming'  # 排除正在流式传输的消息
                ).order_by(ChatMessage.created_at.asc()).all()
                
                logger.info(f"[STREAM_HISTORY_FOUND] 历史消息获取成功 - 总数: {len(history_messages)} 条")
                
                # 准备两种格式的历史记录
                chat_history = []
                for msg in history_messages[-8:]:
                    chat_history.append({
                        "role": "user" if msg.role == "user" else "assistant",
                        "content": msg.content
                    })
                
                context_history = "\n".join([
                    f"{'User' if msg.role == 'user' else 'Assistant'}: {msg.content}"
                    for msg in history_messages[-8:]
                ])
                
                logger.debug(f"[STREAM_CONTEXT] 上下文构建完成 - 结构化历史: {len(chat_history)} 条, 文本历史长度: {len(context_history)} 字符")
            except Exception as e:
                logger.error(f"[STREAM_HISTORY_ERROR] 获取历史消息失败 - 错误: {str(e)}")
                chat_history = []
                context_history = ""
            
            # 5. 准备上下文和生成回答
            full_response = {"answer": "", "sources": [], "confidence": 0.5, "reply": []}
            chunks = []
            update_interval = 5  # 每5个chunk更新一次数据库
            
            # 获取会话相关文档，判断是否有文档上下文
            has_documents = False
            paper_ids = []
            document_info = {"count": 0, "names": []}
            
            if session.session_type != "general":
                logger.info(f"[STREAM_DOCUMENTS] 开始获取会话文档 - 会话ID: {session_id}")
                session_docs = self.db.query(SessionDocument).filter(
                    SessionDocument.session_id == uuid.UUID(session_id)
                ).all()
                
                for doc in session_docs:
                    paper_ids.append(str(doc.paper_id))
                    document_info["names"].append(doc.filename)
                
                has_documents = len(paper_ids) > 0
                document_info["count"] = len(paper_ids)
                logger.info(f"[STREAM_DOCUMENTS_FOUND] 会话文档获取完成 - 文档数量: {len(paper_ids)}")
                
                if paper_ids:
                    logger.debug(f"[STREAM_DOCUMENTS_DETAIL] 文档详情: {document_info}")
            else:
                logger.info(f"[STREAM_DOCUMENTS_SKIP] 通用会话，跳过文档获取")
            
            # 智能分析问题意图 - 使用增强的Gemini意图识别
            logger.info(f"[STREAM_INTENT] 开始Gemini智能意图分析 - 消息: {message[:50]}...")
            intent_result = await self.ai_manager.identify_intent(
                question=message,
                context=None,
                has_documents=has_documents,
                history=context_history,
                document_info=document_info  # 传递详细的文档信息
            )
            
            logger.info(f"[STREAM_INTENT_RESULT] 意图分析完成 - 结果: {intent_result}")
            
            # 6. 根据意图决定处理方式并实现智能路由
            if intent_result == "GENERAL_QUERY":
                # 通用查询模式 - 完全独立的AI对话
                logger.info(f"[STREAM_ROUTE_GENERAL] 路由到通用查询模式 - 会话ID: {session_id}")
                logger.debug(f"[STREAM_ROUTE_GENERAL_DETAIL] 不使用文档上下文，提供独立AI对话")
                
                chunk_count = 0
                async for chunk in self.ai_manager.chat_without_context_stream(
                    message=message,
                    history=chat_history
                ):
                    chunk_count += 1
                    if not chunk.get("error"):
                        text = chunk.get("partial_response", "")
                        full_response["answer"] += text
                        chunks.append(text)
                        
                        logger.debug(f"[STREAM_CHUNK_GENERAL] 通用模式chunk #{chunk_count} - 长度: {len(text)} 字符")
                        
                        # 渐进式保存：每隔几个chunk更新数据库
                        if ai_message and chunk_count % update_interval == 0:
                            try:
                                ai_message.content = full_response["answer"]
                                ai_message.updated_at = datetime.utcnow()
                                self.db.commit()
                                logger.debug(f"[STREAM_PROGRESSIVE_SAVE] 渐进式保存 #{chunk_count//update_interval} - 内容长度: {len(full_response['answer'])} 字符")
                            except Exception as save_error:
                                logger.warning(f"[STREAM_PROGRESSIVE_SAVE_ERROR] 渐进式保存失败: {str(save_error)}")
                                self.db.rollback()
                        
                        yield {
                            "id": f"stream_{len(chunks)}_{session_id}",
                            "content": text,
                            "done": chunk.get("done", False)
                        }
                    else:
                        logger.error(f"[STREAM_ERROR_GENERAL] 通用模式流式错误 - 错误: {chunk.get('error')}")
                        yield {
                            "error": chunk.get("error"),
                            "done": True
                        }
                        return
                
                logger.info(f"[STREAM_COMPLETE_GENERAL] 通用模式流式完成 - 总chunks: {chunk_count}, 响应长度: {len(full_response['answer'])} 字符")
            
            elif intent_result == "DOCUMENT_QUERY":
                # 文档查询模式
                logger.info(f"[STREAM_ROUTE_DOCUMENT] 路由到文档模式 - 会话ID: {session_id}")
                retrieval_context = None

                # 获取文档上下文
                if paper_ids:
                    try:
                        logger.info(f"[STREAM_CONTEXT_RETRIEVAL] 开始文档上下文检索 - 文档数量: {len(paper_ids)}")
                        
                        # 使用标准检索参数
                        doc_limit_per_paper = 4
                        
                        if len(paper_ids) > 1:
                            logger.info(f"[STREAM_CONTEXT_MULTI] 多文档上下文检索 - 每文档限制: {doc_limit_per_paper}")
                            retrieval_context = await self.rag_retriever.get_context_from_multiple_docs(
                                question=message,
                                paper_ids=paper_ids,
                                history=context_history,
                                doc_limit_per_paper=doc_limit_per_paper
                            )
                        else:
                            logger.info(f"[STREAM_CONTEXT_SINGLE] 单文档上下文检索 - 文档ID: {paper_ids[0]}")
                            retrieval_context = await self.rag_retriever.get_relevant_context(
                                question=message,
                                paper_id=paper_ids[0],
                                history=context_history,
                                doc_limit=doc_limit_per_paper * 2
                            )
                        
                        if retrieval_context:
                            context_info = self._analyze_retrieval_context(retrieval_context)
                            logger.info(f"[STREAM_CONTEXT_SUCCESS] 文档上下文获取成功 - {context_info}")
                        else:
                            logger.warning(f"[STREAM_CONTEXT_EMPTY] 未检索到相关文档内容")
                            
                    except Exception as context_error:
                        logger.error(f"[STREAM_CONTEXT_ERROR] 获取文档上下文失败: {str(context_error)}")
                        # 降级为通用模式
                        logger.info(f"[STREAM_FALLBACK_TO_GENERAL] 文档检索失败，降级为通用模式")
                        
                        chunk_count = 0
                        error_context = f"注意：无法访问文档内容（{str(context_error)[:50]}），以下回答基于通用知识。\n\n"
                        full_response["answer"] = error_context
                        
                        async for chunk in self.ai_manager.chat_without_context_stream(
                            message=message,
                            history=chat_history
                        ):
                            chunk_count += 1
                            if not chunk.get("error"):
                                text = chunk.get("partial_response", "")
                                full_response["answer"] += text
                                chunks.append(text)
                                
                                logger.debug(f"[STREAM_FALLBACK_CHUNK] 降级模式chunk #{chunk_count} - 长度: {len(text)} 字符")
                                
                                # 渐进式保存
                                if ai_message and chunk_count % update_interval == 0:
                                    try:
                                        ai_message.content = full_response["answer"]
                                        ai_message.updated_at = datetime.utcnow()
                                        self.db.commit()
                                        logger.debug(f"[STREAM_FALLBACK_PROGRESSIVE_SAVE] 降级模式渐进式保存 #{chunk_count//update_interval}")
                                    except Exception as save_error:
                                        logger.warning(f"[STREAM_FALLBACK_PROGRESSIVE_SAVE_ERROR] 降级模式渐进式保存失败: {str(save_error)}")
                                        self.db.rollback()
                                
                                yield {
                                    "id": f"stream_{len(chunks)}_{session_id}",
                                    "content": text,
                                    "done": chunk.get("done", False)
                                }
                            else:
                                logger.error(f"[STREAM_FALLBACK_ERROR] 降级模式流式错误 - 错误: {chunk.get('error')}")
                                yield {
                                    "error": chunk.get("error"),
                                    "done": True
                                }
                                return
                        
                        logger.info(f"[STREAM_FALLBACK_COMPLETE] 降级模式流式完成 - 总chunks: {chunk_count}")
                        
                        # 确保最终保存
                        if ai_message:
                            confirmation = await self._finalize_ai_message(ai_message, full_response)
                            yield confirmation
                            return
                
                # 使用文档上下文进行流式回答
                if retrieval_context:
                    logger.info(f"[STREAM_AI_DOCUMENT] 开始文档模式流式处理")
                    chunk_count = 0
                    async for chunk in self.ai_manager.stream_response(
                        question=message,
                        context=retrieval_context,
                        history=context_history
                    ):
                        chunk_count += 1
                        if not chunk.get("error"):
                            text = chunk.get("partial_response", "")
                            full_response["answer"] += text
                            chunks.append(text)
                            
                            logger.debug(f"[STREAM_DOC_CHUNK] 文档模式chunk #{chunk_count} - 长度: {len(text)} 字符")
                            
                            # 渐进式保存
                            if ai_message and chunk_count % update_interval == 0:
                                try:
                                    ai_message.content = full_response["answer"]
                                    ai_message.updated_at = datetime.utcnow()
                                    self.db.commit()
                                    logger.debug(f"[STREAM_DOC_PROGRESSIVE_SAVE] 文档模式渐进式保存 #{chunk_count//update_interval}")
                                except Exception as save_error:
                                    logger.warning(f"[STREAM_DOC_PROGRESSIVE_SAVE_ERROR] 文档模式渐进式保存失败: {str(save_error)}")
                                    self.db.rollback()
                            
                            yield {
                                "id": f"stream_{len(chunks)}_{session_id}",
                                "content": text,
                                "done": chunk.get("done", False)
                            }
                        else:
                            logger.error(f"[STREAM_DOC_ERROR] 文档模式流式错误 - 错误: {chunk.get('error')}")
                            yield {
                                "error": chunk.get("error"),
                                "done": True
                            }
                            return
                    
                    logger.info(f"[STREAM_DOC_COMPLETE] 文档模式流式完成 - 总chunks: {chunk_count}, 响应长度: {len(full_response['answer'])} 字符")
                    
                    # 提取来源信息
                    if isinstance(retrieval_context, dict) and "chunks" in retrieval_context:
                        full_response["sources"] = retrieval_context["chunks"][:5]  # 最多5个来源
                        full_response["confidence"] = 0.8  # 基于文档的回答给较高置信度
                        logger.debug(f"[STREAM_DOC_SOURCES] 提取了 {len(full_response['sources'])} 个来源")
            
            elif intent_result in ["SUMMARY_QUERY", "COMPARISON_QUERY", "ANALYSIS_QUERY"]:
                # 特殊文档查询模式（总结、比较、分析）
                logger.info(f"[STREAM_ROUTE_SPECIAL_DOC] 路由到特殊文档模式: {intent_result} - 会话ID: {session_id}")
                retrieval_context = None

                # 获取文档上下文
                if paper_ids:
                    try:
                        logger.info(f"[STREAM_SPECIAL_CONTEXT_RETRIEVAL] 开始特殊文档上下文检索 - 文档数量: {len(paper_ids)}")
                        
                        # 特殊查询使用更多文档内容
                        doc_limit_per_paper = 6 if intent_result == "SUMMARY_QUERY" else 5
                        
                        if len(paper_ids) > 1:
                            logger.info(f"[STREAM_SPECIAL_CONTEXT_MULTI] 多文档特殊上下文检索 - 每文档限制: {doc_limit_per_paper}")
                            retrieval_context = await self.rag_retriever.get_context_from_multiple_docs(
                                question=message,
                                paper_ids=paper_ids,
                                history=context_history,
                                doc_limit_per_paper=doc_limit_per_paper
                            )
                        else:
                            logger.info(f"[STREAM_SPECIAL_CONTEXT_SINGLE] 单文档特殊上下文检索 - 文档ID: {paper_ids[0]}")
                            retrieval_context = await self.rag_retriever.get_relevant_context(
                                question=message,
                                paper_id=paper_ids[0],
                                history=context_history,
                                doc_limit=doc_limit_per_paper * 2
                            )
                        
                        if retrieval_context:
                            context_info = self._analyze_retrieval_context(retrieval_context)
                            logger.info(f"[STREAM_SPECIAL_CONTEXT_SUCCESS] 特殊文档上下文获取成功 - {context_info}")
                        else:
                            logger.warning(f"[STREAM_SPECIAL_CONTEXT_EMPTY] 未检索到相关特殊文档内容")
                            
                    except Exception as context_error:
                        logger.error(f"[STREAM_SPECIAL_CONTEXT_ERROR] 获取特殊文档上下文失败: {str(context_error)}")
                        # 降级为通用模式
                        logger.info(f"[STREAM_SPECIAL_FALLBACK_TO_GENERAL] 特殊文档检索失败，降级为通用模式")
                        
                        chunk_count = 0
                        error_context = f"注意：无法访问文档内容（{str(context_error)[:50]}），以下回答基于通用知识。\n\n"
                        full_response["answer"] = error_context
                        
                        async for chunk in self.ai_manager.chat_without_context_stream(
                            message=message,
                            history=chat_history
                        ):
                            chunk_count += 1
                            if not chunk.get("error"):
                                text = chunk.get("partial_response", "")
                                full_response["answer"] += text
                                chunks.append(text)
                                
                                logger.debug(f"[STREAM_SPECIAL_FALLBACK_CHUNK] 特殊降级模式chunk #{chunk_count} - 长度: {len(text)} 字符")
                                
                                # 渐进式保存
                                if ai_message and chunk_count % update_interval == 0:
                                    try:
                                        ai_message.content = full_response["answer"]
                                        ai_message.updated_at = datetime.utcnow()
                                        self.db.commit()
                                        logger.debug(f"[STREAM_SPECIAL_FALLBACK_PROGRESSIVE_SAVE] 特殊降级模式渐进式保存 #{chunk_count//update_interval}")
                                    except Exception as save_error:
                                        logger.warning(f"[STREAM_SPECIAL_FALLBACK_PROGRESSIVE_SAVE_ERROR] 特殊降级模式渐进式保存失败: {str(save_error)}")
                                        self.db.rollback()
                                
                                yield {
                                    "id": f"stream_{len(chunks)}_{session_id}",
                                    "content": text,
                                    "done": chunk.get("done", False)
                                }
                            else:
                                logger.error(f"[STREAM_SPECIAL_FALLBACK_ERROR] 特殊降级模式流式错误 - 错误: {chunk.get('error')}")
                                yield {
                                    "error": chunk.get("error"),
                                    "done": True
                                }
                                return
                        
                        logger.info(f"[STREAM_SPECIAL_FALLBACK_COMPLETE] 特殊降级模式流式完成 - 总chunks: {chunk_count}")
                
                # 使用文档上下文进行特殊查询的流式回答
                if retrieval_context:
                    logger.info(f"[STREAM_AI_SPECIAL_DOCUMENT] 开始特殊文档模式流式处理: {intent_result}")
                    chunk_count = 0
                    async for chunk in self.ai_manager.stream_response(
                        question=message,
                        context=retrieval_context,
                        history=context_history
                    ):
                        chunk_count += 1
                        if not chunk.get("error"):
                            text = chunk.get("partial_response", "")
                            full_response["answer"] += text
                            chunks.append(text)
                            
                            logger.debug(f"[STREAM_SPECIAL_DOC_CHUNK] 特殊文档模式chunk #{chunk_count} - 长度: {len(text)} 字符")
                            
                            # 渐进式保存
                            if ai_message and chunk_count % update_interval == 0:
                                try:
                                    ai_message.content = full_response["answer"]
                                    ai_message.updated_at = datetime.utcnow()
                                    self.db.commit()
                                    logger.debug(f"[STREAM_SPECIAL_DOC_PROGRESSIVE_SAVE] 特殊文档模式渐进式保存 #{chunk_count//update_interval}")
                                except Exception as save_error:
                                    logger.warning(f"[STREAM_SPECIAL_DOC_PROGRESSIVE_SAVE_ERROR] 特殊文档模式渐进式保存失败: {str(save_error)}")
                                    self.db.rollback()
                            
                            yield {
                                "id": f"stream_{len(chunks)}_{session_id}",
                                "content": text,
                                "done": chunk.get("done", False)
                            }
                        else:
                            logger.error(f"[STREAM_SPECIAL_DOC_ERROR] 特殊文档模式流式错误 - 错误: {chunk.get('error')}")
                            yield {
                                "error": chunk.get("error"),
                                "done": True
                            }
                            return
                    
                    logger.info(f"[STREAM_SPECIAL_DOC_COMPLETE] 特殊文档模式流式完成 - 总chunks: {chunk_count}, 响应长度: {len(full_response['answer'])} 字符")
                    
                    # 提取来源信息
                    if isinstance(retrieval_context, dict) and "chunks" in retrieval_context:
                        full_response["sources"] = retrieval_context["chunks"][:5]  # 最多5个来源
                        full_response["confidence"] = 0.85  # 特殊文档查询给更高置信度
                        logger.debug(f"[STREAM_SPECIAL_DOC_SOURCES] 提取了 {len(full_response['sources'])} 个来源")
            
            else:
                # 未知意图或fallback，使用通用模式
                logger.warning(f"[STREAM_ROUTE_UNKNOWN] 未知意图: {intent_result}, 使用通用模式")
                # ... 通用模式处理逻辑（重复上面的代码）
                async for chunk in self.ai_manager.chat_without_context_stream(
                    message=message,
                    history=chat_history
                ):
                    if not chunk.get("error"):
                        text = chunk.get("partial_response", "")
                        full_response["answer"] += text
                        chunks.append(text)
                        yield {
                            "id": f"stream_{len(chunks)}_{session_id}",
                            "content": text,
                            "done": chunk.get("done", False)
                        }
                    else:
                        yield {"error": chunk.get("error"), "done": True}
                        return
            
            # 7. 最终保存和完成处理
            if ai_message:
                logger.info(f"[STREAM_FINALIZE] 开始最终保存处理 - 消息ID: {ai_message.id}")
                try:
                    # 解析富文本内容
                    try:
                        full_response["reply"] = self.ai_manager._parse_response_content(full_response["answer"])
                        logger.info(f"[STREAM_PARSE_SUCCESS] 富文本解析成功 - 生成 {len(full_response['reply'])} 个内容块")
                    except Exception as parse_error:
                        logger.error(f"[STREAM_PARSE_ERROR] 富文本解析失败，使用fallback - 错误: {str(parse_error)}")
                        full_response["reply"] = [{"type": "markdown", "content": full_response["answer"]}]
                    
                    # 最终更新AI消息
                    ai_message.content = full_response["answer"]
                    ai_message.sources = full_response.get("sources", [])
                    ai_message.confidence = full_response.get("confidence", 0.5)
                    ai_message.rich_content = full_response.get("reply", [])
                    ai_message.message_type = 'markdown'  # 标记流式完成，改为正常消息
                    ai_message.updated_at = datetime.utcnow()
                    
                    # 更新会话时间
                    session.updated_at = datetime.utcnow()
                    
                    self.db.commit()
                    logger.info(f"[STREAM_FINALIZE_SUCCESS] 最终保存成功 - 消息ID: {ai_message.id}")
                    
                    # ✅ 发送保存确认 - 关键修复点
                    confirmation_message = {
                        "message_id": str(ai_message.id),
                        "saved": True,
                        "done": True,
                        "sources": full_response.get("sources", []),
                        "confidence": full_response.get("confidence", 0.0),
                        "reply": full_response.get("reply", [])
                    }
                    logger.info(f"[STREAM_CONFIRMATION_SEND] 发送保存确认消息 - 消息ID: {ai_message.id}")
                    logger.debug(f"[STREAM_CONFIRMATION_CONTENT] 确认消息内容: {confirmation_message}")
                    
                    yield confirmation_message
                    logger.info(f"[STREAM_COMPLETE] 流式处理完成 - 消息ID: {ai_message.id}")
                    
                except Exception as finalize_error:
                    logger.error(f"[STREAM_FINALIZE_ERROR] 最终保存失败: {str(finalize_error)}")
                    self.db.rollback()
                    yield {
                        "error": f"Stream completed but final save failed: {str(finalize_error)}",
                        "done": True
                    }
            else:
                logger.error(f"[STREAM_NO_AI_MESSAGE] 没有AI消息对象，无法保存")
                yield {
                    "error": "No AI message object to save",
                    "done": True
                }
            
        except Exception as e:
            logger.error(f"[STREAM_FATAL_ERROR] 流式处理发生致命错误 - 会话ID: {session_id}, 错误: {str(e)}")
            import traceback
            logger.debug(f"[STREAM_FATAL_TRACEBACK] 致命错误堆栈: {traceback.format_exc()}")
            
            # 即使出错也尝试保存已有内容
            if ai_message and full_response.get("answer"):
                try:
                    logger.info(f"[STREAM_EMERGENCY_SAVE] 尝试紧急保存已有内容")
                    ai_message.content = full_response["answer"] + "\n\n⚠️ 流式传输中断"
                    ai_message.message_type = 'markdown'
                    ai_message.updated_at = datetime.utcnow()
                    self.db.commit()
                    logger.info(f"[STREAM_EMERGENCY_SAVE_SUCCESS] 紧急保存成功")
                    
                    # ✅ 发送紧急保存确认
                    yield {
                        "message_id": str(ai_message.id),
                        "saved": True,
                        "done": True,
                        "emergency_save": True,
                        "sources": full_response.get("sources", []),
                        "confidence": full_response.get("confidence", 0.5),
                        "reply": [{"type": "markdown", "content": full_response["answer"] + "\n\n⚠️ 流式传输中断"}]
                    }
                except Exception as emergency_error:
                    logger.error(f"[STREAM_EMERGENCY_SAVE_ERROR] 紧急保存失败: {str(emergency_error)}")
                    self.db.rollback()
            
            yield {
                "error": str(e),
                "done": True
            }

    async def _finalize_ai_message(self, ai_message: ChatMessage, full_response: dict):
        """最终化AI消息的辅助方法，并返回确认消息"""
        try:
            logger.info(f"[FINALIZE_AI_MESSAGE] 开始最终化AI消息 - 消息ID: {ai_message.id}")
            
            # 解析富文本内容
            try:
                full_response["reply"] = self.ai_manager._parse_response_content(full_response["answer"])
                logger.debug(f"[FINALIZE_PARSE_SUCCESS] 富文本解析成功")
            except Exception as parse_error:
                logger.error(f"[FINALIZE_PARSE_ERROR] 富文本解析失败: {str(parse_error)}")
                full_response["reply"] = [{"type": "markdown", "content": full_response["answer"]}]
            
            # 更新AI消息
            ai_message.content = full_response["answer"]
            ai_message.sources = full_response.get("sources", [])
            ai_message.confidence = full_response.get("confidence", 0.5)
            ai_message.rich_content = full_response.get("reply", [])
            ai_message.message_type = 'markdown'  # 标记流式完成，改为正常消息
            ai_message.updated_at = datetime.utcnow()
            
            self.db.commit()
            logger.info(f"[FINALIZE_AI_MESSAGE_SUCCESS] AI消息最终化成功")
            
            # ✅ 返回确认消息
            return {
                "message_id": str(ai_message.id),
                "saved": True,
                "done": True,
                "sources": full_response.get("sources", []),
                "confidence": full_response.get("confidence", 0.5),
                "reply": full_response.get("reply", [])
            }
            
        except Exception as e:
            logger.error(f"[FINALIZE_AI_MESSAGE_ERROR] AI消息最终化失败: {str(e)}")
            self.db.rollback()
            raise

    async def cleanup_old_sessions(self, days_threshold: int = 7, dry_run: bool = True) -> dict:
        """清理旧的临时会话和无活动会话"""
        logger.info(f"[CLEANUP_SESSIONS_START] 开始清理旧会话 - 阈值: {days_threshold} 天, 模拟运行: {dry_run}")
        
        try:
            from datetime import timedelta
            cutoff_date = datetime.utcnow() - timedelta(days=days_threshold)
            logger.debug(f"[CLEANUP_SESSIONS_CUTOFF] 清理截止日期: {cutoff_date}")
            
            # 查找需要清理的会话
            sessions_to_cleanup = self.db.query(ChatSession).filter(
                (
                    # 临时会话超过阈值天数
                    (ChatSession.is_temporary == True) & 
                    (ChatSession.updated_at < cutoff_date)
                ) | (
                    # 或者无活动的普通会话超过30天
                    (ChatSession.is_temporary == False) & 
                    (ChatSession.updated_at < datetime.utcnow() - timedelta(days=30)) &
                    (ChatSession.message_count == 0)  # 假设有message_count字段
                )
            ).all()
            
            logger.info(f"[CLEANUP_SESSIONS_FOUND] 找到待清理会话 - 数量: {len(sessions_to_cleanup)}")
            
            cleanup_stats = {
                "sessions_found": len(sessions_to_cleanup),
                "messages_found": 0,
                "fragments_found": 0,
                "sessions_cleaned": 0,
                "messages_cleaned": 0,
                "fragments_cleaned": 0,
                "dry_run": dry_run
            }
            
            if not sessions_to_cleanup:
                logger.info(f"[CLEANUP_SESSIONS_EMPTY] 没有找到需要清理的会话")
                return cleanup_stats
            
            # 统计相关消息和分片
            session_ids = [s.id for s in sessions_to_cleanup]
            logger.debug(f"[CLEANUP_SESSIONS_IDS] 待清理会话ID列表: {[str(sid) for sid in session_ids[:5]]}{'...' if len(session_ids) > 5 else ''}")
            
            # 统计消息数量
            messages_to_cleanup = self.db.query(ChatMessage).filter(
                ChatMessage.session_id.in_(session_ids)
            ).all()
            cleanup_stats["messages_found"] = len(messages_to_cleanup)
            logger.info(f"[CLEANUP_SESSIONS_MESSAGES] 找到相关消息 - 数量: {len(messages_to_cleanup)}")
            
            # 统计分片数量
            if messages_to_cleanup:
                from models.chat import MessageFragment
                message_ids = [m.id for m in messages_to_cleanup]
                fragments_count = self.db.query(MessageFragment).filter(
                    MessageFragment.message_id.in_(message_ids)
                ).count()
                cleanup_stats["fragments_found"] = fragments_count
                logger.info(f"[CLEANUP_SESSIONS_FRAGMENTS] 找到相关分片 - 数量: {fragments_count}")
            
            if not dry_run:
                # 执行实际清理
                logger.info(f"[CLEANUP_SESSIONS_EXECUTE] 开始执行实际清理操作")
                try:
                    # 1. 删除消息分片
                    if messages_to_cleanup:
                        from models.chat import MessageFragment
                        message_ids = [m.id for m in messages_to_cleanup]
                        logger.debug(f"[CLEANUP_SESSIONS_DELETE_FRAGMENTS] 删除消息分片")
                        fragments_deleted = self.db.query(MessageFragment).filter(
                            MessageFragment.message_id.in_(message_ids)
                        ).delete(synchronize_session=False)
                        cleanup_stats["fragments_cleaned"] = fragments_deleted
                        logger.info(f"[CLEANUP_SESSIONS_FRAGMENTS_DELETED] 分片删除完成 - 数量: {fragments_deleted}")
                    
                    # 2. 删除消息
                    logger.debug(f"[CLEANUP_SESSIONS_DELETE_MESSAGES] 删除消息")
                    messages_deleted = self.db.query(ChatMessage).filter(
                        ChatMessage.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    cleanup_stats["messages_cleaned"] = messages_deleted
                    logger.info(f"[CLEANUP_SESSIONS_MESSAGES_DELETED] 消息删除完成 - 数量: {messages_deleted}")
                    
                    # 3. 删除会话文档关联
                    logger.debug(f"[CLEANUP_SESSIONS_DELETE_DOCS] 删除会话文档关联")
                    self.db.query(SessionDocument).filter(
                        SessionDocument.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    
                    # 4. 删除会话
                    logger.debug(f"[CLEANUP_SESSIONS_DELETE_SESSIONS] 删除会话")
                    sessions_deleted = self.db.query(ChatSession).filter(
                        ChatSession.id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    cleanup_stats["sessions_cleaned"] = sessions_deleted
                    logger.info(f"[CLEANUP_SESSIONS_SESSIONS_DELETED] 会话删除完成 - 数量: {sessions_deleted}")
                    
                    self.db.commit()
                    logger.info(f"[CLEANUP_SESSIONS_COMMIT] 清理操作提交成功")
                    print(f"清理完成: 删除了{sessions_deleted}个会话, {messages_deleted}个消息, {cleanup_stats['fragments_cleaned']}个分片")
                    
                except Exception as cleanup_error:
                    self.db.rollback()
                    logger.error(f"[CLEANUP_SESSIONS_ERROR] 清理过程中出错: {str(cleanup_error)}")
                    raise Exception(f"清理过程中出错: {str(cleanup_error)}")
            else:
                logger.info(f"[CLEANUP_SESSIONS_DRY_RUN] 模拟清理完成")
                print(f"模拟清理: 将删除{len(sessions_to_cleanup)}个会话, {len(messages_to_cleanup)}个消息, {cleanup_stats['fragments_found']}个分片")
            
            logger.info(f"[CLEANUP_SESSIONS_COMPLETE] 会话清理完成 - 统计: {cleanup_stats}")
            return cleanup_stats
            
        except Exception as e:
            logger.error(f"[CLEANUP_SESSIONS_FATAL] 清理旧会话失败: {str(e)}")
            print(f"清理旧会话失败: {str(e)}")
            raise Exception(f"清理旧会话失败: {str(e)}")

    async def cleanup_orphaned_data(self, dry_run: bool = True) -> dict:
        """清理孤立的数据（没有关联会话的消息等）"""
        try:
            cleanup_stats = {
                "orphaned_messages": 0,
                "orphaned_fragments": 0,
                "orphaned_session_docs": 0,
                "cleaned_messages": 0,
                "cleaned_fragments": 0,
                "cleaned_session_docs": 0,
                "dry_run": dry_run
            }
            
            # 查找孤立的消息（会话不存在）
            orphaned_messages = self.db.query(ChatMessage).filter(
                ~ChatMessage.session_id.in_(
                    self.db.query(ChatSession.id)
                )
            ).all()
            cleanup_stats["orphaned_messages"] = len(orphaned_messages)
            
            # 查找孤立的分片（消息不存在）
            from models.chat import MessageFragment
            orphaned_fragments = self.db.query(MessageFragment).filter(
                ~MessageFragment.message_id.in_(
                    self.db.query(ChatMessage.id)
                )
            ).all()
            cleanup_stats["orphaned_fragments"] = len(orphaned_fragments)
            
            # 查找孤立的会话文档关联（会话不存在）
            orphaned_session_docs = self.db.query(SessionDocument).filter(
                ~SessionDocument.session_id.in_(
                    self.db.query(ChatSession.id)
                )
            ).all()
            cleanup_stats["orphaned_session_docs"] = len(orphaned_session_docs)
            
            if not dry_run:
                # 执行清理
                try:
                    # 清理孤立分片
                    if orphaned_fragments:
                        fragments_deleted = self.db.query(MessageFragment).filter(
                            ~MessageFragment.message_id.in_(
                                self.db.query(ChatMessage.id)
                            )
                        ).delete(synchronize_session=False)
                        cleanup_stats["cleaned_fragments"] = fragments_deleted
                    
                    # 清理孤立消息
                    if orphaned_messages:
                        messages_deleted = self.db.query(ChatMessage).filter(
                            ~ChatMessage.session_id.in_(
                                self.db.query(ChatSession.id)
                            )
                        ).delete(synchronize_session=False)
                        cleanup_stats["cleaned_messages"] = messages_deleted
                    
                    # 清理孤立会话文档关联
                    if orphaned_session_docs:
                        session_docs_deleted = self.db.query(SessionDocument).filter(
                            ~SessionDocument.session_id.in_(
                                self.db.query(ChatSession.id)
                            )
                        ).delete(synchronize_session=False)
                        cleanup_stats["cleaned_session_docs"] = session_docs_deleted
                    
                    self.db.commit()
                    print(f"孤立数据清理完成: 消息{cleanup_stats['cleaned_messages']}, 分片{cleanup_stats['cleaned_fragments']}, 关联{cleanup_stats['cleaned_session_docs']}")
                    
                except Exception as cleanup_error:
                    self.db.rollback()
                    raise Exception(f"孤立数据清理失败: {str(cleanup_error)}")
            
            return cleanup_stats
            
        except Exception as e:
            print(f"孤立数据清理失败: {str(e)}")
            raise Exception(f"孤立数据清理失败: {str(e)}")

    # 需要添加类似代码:
    app = Celery('paper_analyzer')
    
    @app.task
    def build_index_task(content, paper_id, line_mapping=None):
        # 异步构建索引
        # ...

    # 在关键操作处添加详细日志
        logger.info(f"Processing document {paper_id} with size {len(content)}")

    def _analyze_retrieval_context(self, retrieval_context) -> str:
        """分析检索到的上下文信息，返回摘要信息"""
        try:
            if isinstance(retrieval_context, dict):
                chunks = retrieval_context.get("chunks", [])
                total_chunks = retrieval_context.get("total_chunks", len(chunks))
                context_text = retrieval_context.get("text", "")
                
                info_parts = []
                info_parts.append(f"检索到 {len(chunks)} 个相关片段（总共 {total_chunks} 个可用）")
                
                if context_text:
                    info_parts.append(f"上下文长度: {len(context_text)} 字符")
                
                # 分析来源多样性
                if chunks:
                    sources = set()
                    for chunk in chunks:
                        if isinstance(chunk, dict) and "document_id" in chunk:
                            sources.add(chunk["document_id"])
                        elif isinstance(chunk, dict) and "document_name" in chunk:
                            sources.add(chunk["document_name"])
                    
                    if sources:
                        info_parts.append(f"涉及 {len(sources)} 个文档来源")
                    
                    # 分析相似度
                    similarities = []
                    for chunk in chunks:
                        if isinstance(chunk, dict) and "similarity" in chunk:
                            similarities.append(chunk["similarity"])
                    
                    if similarities:
                        avg_sim = sum(similarities) / len(similarities)
                        max_sim = max(similarities)
                        info_parts.append(f"相似度: 最高 {max_sim:.2f}, 平均 {avg_sim:.2f}")
                
                return ", ".join(info_parts)
            else:
                return f"上下文类型: {type(retrieval_context)}, 长度: {len(str(retrieval_context))} 字符"
                
        except Exception as e:
            logger.error(f"[CONTEXT_ANALYSIS_ERROR] 分析上下文信息失败: {str(e)}")
            return f"上下文分析失败: {str(e)}"

    async def analyze_web_page(self, url: str, user_id: str = None) -> dict:
        """分析网页内容 - 新增方法"""
        start_time = time.time()
        try:
            logger.info(f"[ANALYZE_WEB_START] 开始分析网页 - URL: {url}")
            
            # 1. 处理网页内容
            processed_result = await self.paper_processor.process_url(url)
            
            # 2. 生成唯一的web_id
            web_id = str(uuid.uuid4())
            
            # 3. 保存到数据库
            from models.paper import PaperAnalysis
            web_analysis = PaperAnalysis(
                paper_id=uuid.UUID(web_id),
                filename=processed_result.get("title", url),
                content=processed_result["content"],
                line_mapping=processed_result["line_mapping"],
                total_lines=processed_result["total_lines"],
                structured_data=processed_result.get("structured_data"),
                file_type="web",
                source_url=url,
                documents=None,
                embeddings=None,
                index_built=False,
                user_id=uuid.UUID(user_id) if user_id else None
            )
            self.db.add(web_analysis)
            self.db.commit()
            
            # 4. 异步构建向量索引
            try:
                if self.config.enable_async_indexing:
                    # 构建网页索引
                    success = await self.rag_retriever.build_web_index(
                        processed_result, web_id, self.db
                    )
                    if success:
                        web_analysis.index_built = True
                        self.db.commit()
                        logger.info(f"[ANALYZE_WEB_INDEX] 网页索引构建成功")
                else:
                    # 同步构建索引
                    success = await self.rag_retriever.build_web_index(
                        processed_result, web_id, self.db
                    )
                    web_analysis.index_built = success
                    self.db.commit()
                    
            except Exception as e:
                logger.error(f"[ANALYZE_WEB_INDEX_ERROR] 网页索引构建失败: {str(e)}")
                # 不中断流程，允许用户查看网页内容
            
            # 5. 记忆存储 - 存储网页信息到AI记忆中
            if self.ai_manager.memory_manager and user_id:
                web_summary = processed_result["content"][:500]  # 网页摘要
                self.ai_manager.memory_manager.add_memory(
                    content=f"用户分析网页: {url}，标题: {processed_result.get('title', '未知')}，内容摘要: {web_summary}",
                    context_type="web",
                    importance=0.7,
                    user_id=user_id,
                    document_id=web_id
                )
                logger.debug(f"[ANALYZE_WEB_MEMORY] 已存储网页到记忆系统")
            
            # 6. 更新性能统计
            processing_time = time.time() - start_time
            self.performance_stats["total_documents_processed"] += 1
            self.performance_stats["average_response_time"] = (
                self.performance_stats["average_response_time"] * 0.9 + processing_time * 0.1
            )
            
            logger.info(f"[ANALYZE_WEB_SUCCESS] 网页分析完成 - 耗时: {processing_time:.2f}s")
            
            return {
                "status": "success",
                "message": "网页分析完成",
                "web_id": web_id,
                "url": url,
                "title": processed_result.get("title", ""),
                "content": processed_result["content"],
                "line_mapping": processed_result["line_mapping"] or {},
                "total_lines": processed_result["total_lines"],
                "structured_data": processed_result.get("structured_data", {}),
                "metadata": processed_result.get("metadata", {}),
                "has_index": web_analysis.index_built
            }
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"[ANALYZE_WEB_ERROR] 网页分析失败: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

    async def analyze_multiple_web_pages(self, urls: List[str], user_id: str = None) -> List[dict]:
        """批量分析多个网页 - 新增方法"""
        try:
            logger.info(f"[ANALYZE_WEBS_BATCH] 开始批量分析网页 - 数量: {len(urls)}")
            
            # 批量处理网页
            processed_results = await self.paper_processor.process_multiple_urls(urls)
            
            results = []
            for i, processed_result in enumerate(processed_results):
                try:
                    # 生成web_id
                    web_id = str(uuid.uuid4())
                    url = urls[i] if i < len(urls) else processed_result.get("source_url", "")
                    
                    # 保存到数据库
                    from models.paper import PaperAnalysis
                    web_analysis = PaperAnalysis(
                        paper_id=uuid.UUID(web_id),
                        filename=processed_result.get("title", url),
                        content=processed_result["content"],
                        line_mapping=processed_result["line_mapping"],
                        total_lines=processed_result["total_lines"],
                        structured_data=processed_result.get("structured_data"),
                        file_type="web",
                        source_url=url,
                        documents=None,
                        embeddings=None,
                        index_built=False,
                        user_id=uuid.UUID(user_id) if user_id else None
                    )
                    self.db.add(web_analysis)
                    
                    # 异步构建索引
                    try:
                        success = await self.rag_retriever.build_web_index(
                            processed_result, web_id, self.db
                        )
                        web_analysis.index_built = success
                    except Exception as index_error:
                        logger.error(f"[ANALYZE_WEBS_INDEX_ERROR] 网页索引构建失败 {url}: {str(index_error)}")
                    
                    results.append({
                        "status": "success",
                        "web_id": web_id,
                        "url": url,
                        "title": processed_result.get("title", ""),
                        "has_index": web_analysis.index_built
                    })
                    
                except Exception as single_error:
                    logger.error(f"[ANALYZE_WEBS_SINGLE_ERROR] 单个网页处理失败 {urls[i] if i < len(urls) else 'unknown'}: {str(single_error)}")
                    results.append({
                        "status": "error",
                        "url": urls[i] if i < len(urls) else "unknown",
                        "message": str(single_error)
                    })
            
            self.db.commit()
            logger.info(f"[ANALYZE_WEBS_BATCH_SUCCESS] 批量网页分析完成 - 成功: {len([r for r in results if r['status'] == 'success'])}/{len(urls)}")
            
            return results
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"[ANALYZE_WEBS_BATCH_ERROR] 批量网页分析失败: {str(e)}")
            return [{"status": "error", "message": str(e)} for _ in urls]

    async def search_web_content(self, query: str, user_id: str = None, limit: int = 10) -> dict:
        """搜索网页内容 - 新增方法"""
        try:
            logger.info(f"[SEARCH_WEB] 开始搜索网页内容 - 查询: {query[:50]}...")
            
            # 使用RAG检索器搜索网页
            search_results = await self.rag_retriever.search_web_content(
                query=query,
                user_id=user_id,
                limit=limit
            )
            
            logger.info(f"[SEARCH_WEB_SUCCESS] 网页搜索完成 - 找到结果: {len(search_results.get('results', []))}")
            return search_results
            
        except Exception as e:
            logger.error(f"[SEARCH_WEB_ERROR] 网页搜索失败: {str(e)}")
            return {
                "status": "error",
                "message": str(e),
                "results": [],
                "total": 0
            }

    async def ask_web_question(self, question: str, web_id: str, session_id: str = None, user_id: str = None) -> dict:
        """向网页内容提问 - 新增方法"""
        try:
            logger.info(f"[ASK_WEB] 开始向网页提问 - 网页ID: {web_id}, 问题: {question[:50]}...")
            
            # 获取网页上下文
            web_context = await self.rag_retriever.get_web_context(
                question=question,
                web_id=web_id
            )
            
            if not web_context or not web_context.get("chunks"):
                return {
                    "status": "error",
                    "message": "无法获取网页内容或网页索引未构建"
                }
            
            # 使用AI分析网页内容
            response = await self.ai_manager.analyze_web_content(
                question=question,
                web_context=web_context
            )
            
            # 记忆存储 - 存储问答到记忆系统
            if self.ai_manager.memory_manager and user_id:
                self.ai_manager.memory_manager.add_memory(
                    content=f"网页问答: Q: {question} A: {response['answer'][:200]}",
                    context_type="conversation",
                    importance=0.6,
                    user_id=user_id,
                    document_id=web_id,
                    session_id=session_id
                )
            
            logger.info(f"[ASK_WEB_SUCCESS] 网页问答完成")
            return {
                "status": "success",
                "answer": response["answer"],
                "sources": response.get("sources", []),
                "confidence": response.get("confidence", 0.0),
                "reply": response.get("reply", []),
                "web_info": response.get("web_info", {}),
                "analysis_type": response.get("analysis_type", "web_content")
            }
            
        except Exception as e:
            logger.error(f"[ASK_WEB_ERROR] 网页问答失败: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

    async def create_web_chat_session(self, user_id: str, web_ids: List[str], title: str = None) -> dict:
        """创建网页对话会话 - 新增方法"""
        try:
            logger.info(f"[CREATE_WEB_SESSION] 创建网页对话会话 - 用户ID: {user_id}, 网页数: {len(web_ids)}")
            
            # 验证网页是否存在
            from models.paper import PaperAnalysis
            web_records = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.paper_id.in_([uuid.UUID(wid) for wid in web_ids]),
                PaperAnalysis.file_type == "web"
            ).all()
            
            if len(web_records) != len(web_ids):
                raise ValueError("部分网页ID不存在或不是网页类型")
            
            # 生成标题
            if not title:
                if len(web_records) == 1:
                    title = f"关于 {web_records[0].filename} 的网页对话"
                else:
                    title = f"多网页对话 ({len(web_records)}个网页)"
            
            # 创建会话
            session_result = await self.create_chat_session(
                user_id=user_id,
                title=title,
                paper_ids=web_ids,
                session_type="web"
            )
            
            logger.info(f"[CREATE_WEB_SESSION_SUCCESS] 网页对话会话创建成功 - 会话ID: {session_result['id']}")
            return session_result
            
        except Exception as e:
            logger.error(f"[CREATE_WEB_SESSION_ERROR] 创建网页对话会话失败: {str(e)}")
            raise Exception(f"创建网页对话会话失败: {str(e)}")

    async def get_user_web_pages(self, user_id: str, limit: int = 20) -> List[dict]:
        """获取用户的网页列表 - 新增方法"""
        try:
            logger.info(f"[GET_USER_WEBS] 获取用户网页列表 - 用户ID: {user_id}")
            
            from models.paper import PaperAnalysis
            web_records = self.db.query(PaperAnalysis).filter(
                PaperAnalysis.user_id == uuid.UUID(user_id),
                PaperAnalysis.file_type == "web"
            ).order_by(PaperAnalysis.created_at.desc()).limit(limit).all()
            
            result = []
            for record in web_records:
                result.append({
                    "web_id": str(record.paper_id),
                    "title": record.filename,
                    "url": record.source_url,
                    "content_preview": record.content[:200] if record.content else "",
                    "created_at": record.created_at.isoformat() if record.created_at else None,
                    "updated_at": record.updated_at.isoformat() if record.updated_at else None,
                    "has_index": record.index_built,
                    "total_lines": record.total_lines,
                    "summary": record.summary
                })
            
            logger.info(f"[GET_USER_WEBS_SUCCESS] 获取用户网页列表成功 - 数量: {len(result)}")
            return result
            
        except Exception as e:
            logger.error(f"[GET_USER_WEBS_ERROR] 获取用户网页列表失败: {str(e)}")
            raise Exception(f"获取用户网页列表失败: {str(e)}")

    async def compare_web_and_document(self, question: str, web_id: str, document_id: str, 
                                      user_id: str = None) -> dict:
        """对比网页和文档内容 - 新增方法"""
        try:
            logger.info(f"[COMPARE_WEB_DOC] 开始对比网页和文档 - 网页ID: {web_id}, 文档ID: {document_id}")
            
            # 获取网页上下文
            web_context = await self.rag_retriever.get_web_context(
                question=question,
                web_id=web_id
            )
            
            # 获取文档上下文
            doc_context = await self.rag_retriever.get_relevant_context(
                question=question,
                paper_id=document_id
            )
            
            if not web_context or not doc_context:
                raise Exception("无法获取网页或文档内容")
            
            # 使用AI进行对比分析
            response = await self.ai_manager.analyze_web_vs_document(
                question=question,
                web_context=web_context,
                doc_context=doc_context
            )
            
            # 记忆存储
            if self.ai_manager.memory_manager and user_id:
                self.ai_manager.memory_manager.add_memory(
                    content=f"网页文档对比: {question} - 网页: {web_id}, 文档: {document_id}",
                    context_type="comparison",
                    importance=0.8,
                    user_id=user_id
                )
            
            logger.info(f"[COMPARE_WEB_DOC_SUCCESS] 网页文档对比完成")
            return {
                "status": "success",
                "answer": response["answer"],
                "sources": response.get("sources", []),
                "confidence": response.get("confidence", 0.0),
                "reply": response.get("reply", []),
                "web_info": response.get("web_info", {}),
                "doc_info": response.get("doc_info", {}),
                "analysis_type": response.get("analysis_type", "web_document_comparison")
            }
            
        except Exception as e:
            logger.error(f"[COMPARE_WEB_DOC_ERROR] 网页文档对比失败: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
